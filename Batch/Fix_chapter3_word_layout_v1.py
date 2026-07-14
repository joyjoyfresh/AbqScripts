from __future__ import print_function

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / 'docs' / '论文材料' / '边坡地震动放大效应研究论文初稿（第三章重构）.docx'
FIG = ROOT / 'Run' / 'ch3_F1_frequency_theory' / 'run-003' / 'F1_A1_transfer_function.png'


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def main():
    doc = Document(str(DOCX))

    # 将仅适用于公式编辑器的 LaTeX 片段改为 Word 正文可直接阅读的文本。
    replacements = {
        r'\(\rho\ddot{\boldsymbol u}=\nabla\cdot\boldsymbol\sigma+\boldsymbol b\)': 'ρ ü = div σ + b',
        r'\(\boldsymbol\sigma=\boldsymbol C: \boldsymbol\varepsilon\)': 'σ = C:ε',
        r'\(\boldsymbol u\)': 'u',
        r'\(p=\sin\theta/V_s\)': 'p = sin(θ) / Vs',
        r'\(H_q(f)=U_q^{\rm surface}(f)/U^{\rm input}(f)\)': 'Hq(f) = Uq(surface)(f) / U(input)(f)',
        r'\([0.5f_c,1.5f_c]\)': '[0.5fc, 1.5fc]',
        r'\(\sqrt{|H_h|^2+|H_v|^2}\)': 'sqrt(|Hh|² + |Hv|²)',
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            paragraph.text = new_text

    # 删除先前插入但在分页处被截断的图形段落。
    for paragraph in list(doc.paragraphs):
        if paragraph._p.xpath('.//w:drawing'):
            remove_paragraph(paragraph)
    # 清理多次迭代留下的旧图题，避免重复显示。
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip().startswith('图3-1'):
            remove_paragraph(paragraph)

    # 在 A1 结果段落后重新插入尺寸受控的图和图题。
    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith('3.5.3 平坦成层场地独立一维对比'):
            anchor = paragraph
            break
    if anchor is not None and FIG.exists():
        image_paragraph = anchor.insert_paragraph_before()
        image_paragraph.alignment = 1
        image_paragraph.paragraph_format.keep_together = True
        image_paragraph.paragraph_format.space_before = Pt(3)
        image_paragraph.paragraph_format.space_after = Pt(3)
        image_paragraph.add_run().add_picture(str(FIG), width=Inches(4.9))
        caption = anchor.insert_paragraph_before('图3-1  A1均质半空间三个内部测点的传递函数比较')
        caption.alignment = 1
        caption.paragraph_format.keep_with_next = False
        caption.paragraph_format.space_after = Pt(6)

    doc.save(str(DOCX))
    print('已修正公式文本并重新插入 A1 图形：{}'.format(DOCX))


if __name__ == '__main__':
    main()
