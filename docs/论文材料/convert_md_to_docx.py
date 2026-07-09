import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re
import os

def set_run_font(run, size_pt=10.5, bold=False, italic=False, color_rgb=None, font_name='Calibri'):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color_rgb:
        run.font.color.rgb = color_rgb
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_styled_text_to_p(p, text, is_italic=False, size_pt=10.5, color_rgb=None, font_name='Calibri'):
    # 解析 markdown 中的 **加粗** 和 [文字](链接) 形式（直接过滤链接保留文字）
    # 1. 过滤链接，例如 [utils.py](file://...) -> utils.py
    text_clean = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    
    # 2. 识别 **加粗**
    parts = re.split(r'(\*\*.*?\*\*)', text_clean)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            run = p.add_run(content)
            set_run_font(run, size_pt=size_pt, bold=True, italic=is_italic, color_rgb=color_rgb, font_name=font_name)
        else:
            if part:
                run = p.add_run(part)
                set_run_font(run, size_pt=size_pt, bold=False, italic=is_italic, color_rgb=color_rgb, font_name=font_name)

def add_heading(doc, text, level):
    p = doc.add_heading(level=level)
    p.paragraph_format.keep_with_next = True
    
    # 调整不同级别标题的字号和间距
    if level == 1:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        size_pt = 15
    elif level == 2:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        size_pt = 13
    else:
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        size_pt = 11.5
        
    # 标题为粗体、主题深蓝色 (RGB: 31, 78, 121)
    # 标题可能含有 **加粗**，我们把标题中多余的 ** 剥掉，统一由标题样式加粗
    clean_text = text.replace("**", "")
    run = p.add_run(clean_text)
    set_run_font(run, size_pt=size_pt, bold=True, color_rgb=RGBColor(31, 78, 121))
    return p

def create_code_block(doc, code_lines):
    # 用 1x1 浅灰色底色表格包裹代码，使用 Consolas 字体
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    
    # 设置浅灰色填充底色
    shading_elm = parse_xml(r'<w:shd {} w:fill="F4F4F4"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    
    code_text = "\n".join(code_lines)
    run = p.add_run(code_text)
    set_run_font(run, size_pt=9.0, font_name='Consolas', color_rgb=RGBColor(50, 50, 50))
    
    # 段后加点空白
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(2)
    p_space.paragraph_format.space_after = Pt(2)

def create_blockquote(doc, quote_lines):
    # 用 1x1 浅灰色底色表格包裹引用，使用小五号斜体
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    
    # 设置淡灰色底色
    shading_elm = parse_xml(r'<w:shd {} w:fill="FAFAFA"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    quote_text = "\n".join(quote_lines)
    add_styled_text_to_p(p, quote_text, is_italic=True, size_pt=9.5, color_rgb=RGBColor(80, 80, 80))
    
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(2)
    p_space.paragraph_format.space_after = Pt(2)

def create_table_from_markdown(doc, table_lines):
    # 解析 markdown 表格行
    raw_rows = []
    for line in table_lines:
        # 移除行首尾的 |
        content = line.strip()
        if content.startswith('|'):
            content = content[1:]
        if content.endswith('|'):
            content = content[:-1]
        
        cells = [c.strip() for c in content.split('|')]
        raw_rows.append(cells)
        
    if not raw_rows:
        return
        
    # 过滤掉分隔符行 (如 |--|--|)
    parsed_rows = []
    for row in raw_rows:
        # 如果整行都是减号/冒号/空格，代表是分隔符行
        is_separator = all(re.match(r'^[\s\-\:\=\|]+$', cell) for cell in row)
        if not is_separator:
            parsed_rows.append(row)
            
    if not parsed_rows:
        return
        
    num_cols = max(len(row) for row in parsed_rows)
    num_rows = len(parsed_rows)
    
    table = doc.add_table(rows=num_rows, cols=num_cols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 填充表头
    hdr_cells = table.rows[0].cells
    headers = parsed_rows[0]
    for i in range(num_cols):
        text = headers[i] if i < len(headers) else ""
        hdr_cells[i].text = text
        # 设置底色
        shading_elm = parse_xml(r'<w:shd {} w:fill="1F4E79"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
        
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if p.runs:
            set_run_font(p.runs[0], size_pt=9.5, bold=True, color_rgb=RGBColor(255, 255, 255))
            
    # 填充数据行
    for r_idx in range(1, num_rows):
        row_data = parsed_rows[r_idx]
        row_cells = table.rows[r_idx].cells
        bg_color = "F9F9F9" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx in range(num_cols):
            text = row_data[c_idx] if c_idx < len(row_data) else ""
            row_cells[c_idx].text = text
            if bg_color != "FFFFFF":
                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
                row_cells[c_idx]._tc.get_or_add_tcPr().append(shading_elm)
                
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            # 对齐
            if num_cols > 2 and c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            if p.runs:
                # 重新应用字体
                # 将文本清空并重新用 add_styled_text_to_p 填充以保持加粗效果
                p.text = ""
                add_styled_text_to_p(p, text, size_pt=9.0)
                
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(3)
    p_space.paragraph_format.space_after = Pt(3)

def convert_md_to_docx(md_path=None, docx_path=None):
    if md_path is None:
        md_path = r"c:\Users\12462\Documents\Code\AbqScripts\docs\论文材料\论文章节_模型尺寸设计_v1.md"
    if docx_path is None:
        docx_path = r"c:\Users\12462\Documents\Code\AbqScripts\docs\论文材料\边坡地震响应数值模拟尺寸设计.docx"
    
    if not os.path.exists(md_path):
        print(f"Error: Source Markdown file not found: {md_path}")
        return
        
    doc = Document()
    
    # A4 边距
    for section in doc.sections:
        section.top_margin = docx.shared.Cm(2.54)
        section.bottom_margin = docx.shared.Cm(2.54)
        section.left_margin = docx.shared.Cm(2.54)
        section.right_margin = docx.shared.Cm(2.54)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    in_code_block = False
    code_block_lines = []
    
    in_table = False
    table_lines = []
    
    in_blockquote = False
    blockquote_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        line_strip = line.strip()
        
        # 1. 代码块处理
        if line_strip.startswith("```"):
            if in_code_block:
                # 代码块结束，生成它
                create_code_block(doc, code_block_lines)
                in_code_block = False
                code_block_lines = []
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_block_lines.append(line.rstrip('\n'))
            i += 1
            continue
            
        # 2. 表格处理
        # 判定是否属于表格行：带有 '|' 且非首尾空行
        is_table_line = '|' in line and (line_strip.startswith('|') or line_strip.endswith('|'))
        if is_table_line:
            in_table = True
            table_lines.append(line_strip)
            i += 1
            continue
        elif in_table:
            # 表格行结束，生成它
            create_table_from_markdown(doc, table_lines)
            in_table = False
            table_lines = []
            # 注意：此处不增加 i，让当前行被其他规则处理
            
        # 3. 引用块处理
        if line_strip.startswith(">"):
            in_blockquote = True
            # 剥离 '>' 和多余的空格
            content = line_strip[1:].strip()
            blockquote_lines.append(content)
            i += 1
            continue
        elif in_blockquote:
            create_blockquote(doc, blockquote_lines)
            in_blockquote = False
            blockquote_lines = []
            
        # 4. 标题处理
        if line_strip.startswith("#"):
            # 统计 # 个数
            level = 0
            for char in line_strip:
                if char == '#':
                    level += 1
                else:
                    break
            heading_text = line_strip[level:].strip()
            add_heading(doc, heading_text, level)
            i += 1
            continue
            
        # 5. 无序列表处理
        is_bullet = line_strip.startswith("- ") or line_strip.startswith("* ")
        if is_bullet:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1.5)
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.line_spacing = 1.15
            content = line_strip[2:].strip()
            add_styled_text_to_p(p, content)
            i += 1
            continue
            
        # 6. 有序列表处理
        match_ordered = re.match(r'^(\d+)\.\s+(.*)$', line_strip)
        if match_ordered:
            # 转化为列表项目
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1.5)
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.line_spacing = 1.15
            
            num = match_ordered.group(1)
            content = match_ordered.group(2)
            
            # 把序号写在前面加粗
            run_num = p.add_run(f"{num}. ")
            set_run_font(run_num, bold=True)
            
            add_styled_text_to_p(p, content)
            i += 1
            continue
            
        # 7. 图片处理
        match_image = re.match(r'^!\[(.*?)\]\((.*?)\)$', line_strip)
        if match_image:
            caption = match_image.group(1)
            img_path = match_image.group(2)
            img_path = img_path.replace("%20", " ")
            base_dir = os.path.dirname(os.path.abspath(__file__))
            full_img_path = os.path.join(base_dir, img_path)
            
            if os.path.exists(full_img_path):
                try:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_img = p_img.add_run()
                    run_img.add_picture(full_img_path, width=docx.shared.Cm(14.0))
                    
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_before = Pt(2)
                    p_cap.paragraph_format.space_after = Pt(6)
                    run_cap = p_cap.add_run(caption)
                    set_run_font(run_cap, size_pt=9.0, italic=True, color_rgb=RGBColor(80, 80, 80))
                except Exception as e:
                    print(f"Warning: Failed to insert image {full_img_path}: {e}")
            else:
                print(f"Warning: Image file not found: {full_img_path}")
            i += 1
            continue
            
        # 8. 普通段落与空行
        if not line_strip:
            # 空行略过
            i += 1
            continue
            
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.25
        add_styled_text_to_p(p, line_strip)
        i += 1
        
    # 处理文件尾部未闭合的状态
    if in_table:
        create_table_from_markdown(doc, table_lines)
    if in_blockquote:
        create_blockquote(doc, blockquote_lines)
        
    doc.save(docx_path)
    print(f"Markdown successfully converted to Word at: {docx_path}")

if __name__ == "__main__":
    convert_md_to_docx()
