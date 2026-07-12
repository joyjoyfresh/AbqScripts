# -*- coding: utf-8 -*-
"""将通过 U3a 验证的网格收敛内容写入论文修订稿第3.5.1节。"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


DOCX_PATH = r"C:\Users\12462\Documents\Code\AbqScripts\docs\论文材料\边坡地震动放大效应研究论文初稿（整合修订）.docx"


def set_style_id(paragraph, style_id):
    """直接写入稳定的 Word 样式 ID，避免重复本地化样式名歧义。"""
    p_pr = paragraph._p.get_or_add_pPr()
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        p_style = OxmlElement("w:pStyle")
        p_pr.insert(0, p_style)
    p_style.set(qn("w:val"), style_id)


def insert_paragraph_after(document, anchor, text, style_id, alignment=None):
    """在指定 XML 元素后插入段落。"""
    element = OxmlElement("w:p")
    anchor.addnext(element)
    from docx.text.paragraph import Paragraph
    paragraph = Paragraph(element, document._body)
    set_style_id(paragraph, style_id)
    paragraph.add_run(text)
    if alignment is not None:
        paragraph.alignment = alignment
    return paragraph


def configure_table(table):
    """添加边框、禁止跨页拆行并重复表头。"""
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        item = OxmlElement("w:" + edge)
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), "4")
        item.set(qn("w:color"), "000000")
        borders.append(item)
    table._tbl.tblPr.append(borders)
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))


def main():
    document = Document(DOCX_PATH)
    start = next(p for p in document.paragraphs if p.text.strip().startswith("3.5 "))
    end = next(p for p in document.paragraphs if p.text.strip().startswith("3.6 "))
    start.text = "3.5 数值离散与参数收敛"

    node = start._p.getnext()
    while node is not None and node is not end._p:
        following = node.getnext()
        node.getparent().remove(node)
        node = following

    anchor = start._p
    heading = insert_paragraph_after(document, anchor, "3.5.1 网格收敛", "Heading3")
    anchor = heading._p
    texts = [
        "空间离散采用四节点缩减积分平面应变单元 CPE4R。网格尺寸首先满足最短有效剪切波长的解析要求，再通过目标响应收敛试验确定。U3a 固定均质基岩、坡高25 m、坡角30°、垂直入射、侧向净空6h、坡脚以下深度5h、1%阻尼、4 Hz Ricker 波和0.001 s最大时间增量，只将全局单元尺寸设为12 m、8 m和6 m。",
        "为避免单个极值偶然重合造成假收敛，本文同时采用整曲线和峰值两项指标。在固定501点归一化坐标上，以6 m网格为参考，计算8 m网格 TAF_h 曲线的 L2 相对差，并比较两者最大放大系数 AR_max 的相对差；两项均不超过5%时认为网格收敛。",
        "12 m、8 m和6 m网格的 AR_max 分别为1.081579、1.079491和1.079491。8 m相对6 m的整曲线 L2 相对差为0.134%，峰值相对差为0.000018%，均远低于5%门槛。由此，均质基岩基准工况采用8 m网格即可获得网格无关的目标响应；继续细化至6 m没有带来可辨识的精度收益。",
        "8 m仅作为后续均质基岩工况的基准尺寸。对于存在低波速薄层的成层坡地，脚本仍以控制性最软层的最短有效波长、厚度方向穿层单元数和单元尺寸下限共同约束局部网格，不能直接套用8 m的绝对尺寸。时间步、计算域和阻尼的最终取值将在后续独立收敛试验通过后确定。"
    ]
    for text in texts:
        paragraph = insert_paragraph_after(document, anchor, text, "BodyText", WD_ALIGN_PARAGRAPH.JUSTIFY)
        anchor = paragraph._p

    caption = insert_paragraph_after(document, anchor, "表3.3 网格收敛结果", "FigCaption", WD_ALIGN_PARAGRAPH.CENTER)
    anchor = caption._p
    table = document.add_table(rows=4, cols=4)
    headers = ["网格尺寸", "AR_max", "相对6 m峰值差", "说明"]
    rows = [
        ["12 m", "1.081579", "0.1935%", "粗网格参考"],
        ["8 m", "1.079491", "0.000018%", "后续基准网格"],
        ["6 m", "1.079491", "—", "收敛参考"]
    ]
    for column, text in enumerate(headers):
        table.cell(0, column).text = text
    for row_index, values in enumerate(rows, 1):
        for column, text in enumerate(values):
            table.cell(row_index, column).text = text
    configure_table(table)
    anchor.addnext(table._tbl)

    document.save(DOCX_PATH)
    print("已写入第3.5.1节并保存：%s" % DOCX_PATH)


if __name__ == "__main__":
    main()
