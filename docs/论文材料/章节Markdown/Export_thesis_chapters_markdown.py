from __future__ import print_function

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[3]
SOURCE = ROOT / 'docs' / '论文材料' / '边坡地震动放大效应研究论文初稿（第三章重构）.docx'
OUTPUT = ROOT / 'docs' / '论文材料' / '章节Markdown'


SECTION_FILES = {
    '摘    要': ('00_摘要.md', '摘要'),
    'Abstract': ('00_Abstract.md', 'Abstract'),
    '目    录': ('00_目录.md', '目录'),
    '1 绪论': ('第1章_绪论.md', '第1章 绪论'),
    '2 坡地地震波传播理论与波动输入方法': ('第2章_坡地地震波传播理论与波动输入方法.md', '第2章 坡地地震波传播理论与波动输入方法'),
    '3 斜入射成层坡地有限元模型与可信性验证': ('第3章_斜入射成层坡地有限元模型与可信性验证.md', '第3章 斜入射成层坡地有限元模型与可信性验证'),
    '4 坡地地震动放大效应参数化分析': ('第4章_坡地地震动放大效应参数化分析.md', '第4章 坡地地震动放大效应参数化分析'),
    '5 地形-土-结构相互作用分析': ('第5章_地形-土-结构相互作用分析.md', '第5章 地形-土-结构相互作用分析'),
    '6 基于机器学习的坡地地震动放大效应预测模型': ('第6章_基于机器学习的坡地地震动放大效应预测模型.md', '第6章 基于机器学习的坡地地震动放大效应预测模型'),
    '7 结论与展望': ('第7章_结论与展望.md', '第7章 结论与展望'),
    '参考文献': ('参考文献.md', '参考文献'),
    '附    录': ('附录.md', '附录'),
}

NORMALIZED_SECTION_FILES = {re.sub(r'\s+', ' ', key).strip(): value for key, value in SECTION_FILES.items()}


def clean_text(text):
    text = text.replace('\u00a0', ' ')
    text = re.sub(r'[ \t]+', ' ', text)
    return text.strip()


def markdown_table(table):
    rows = []
    for row in table.rows:
        cells = [clean_text(cell.text).replace('|', '\\|').replace('\n', '<br>') for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ''
    width = max(len(row) for row in rows)
    rows = [row + [''] * (width - len(row)) for row in rows]
    lines = ['| ' + ' | '.join(rows[0]) + ' |',
             '| ' + ' | '.join(['---'] * width) + ' |']
    lines.extend('| ' + ' | '.join(row) + ' |' for row in rows[1:])
    return '\n'.join(lines)


def extract_images(doc):
    attachments = OUTPUT / '附件'
    attachments.mkdir(parents=True, exist_ok=True)
    image_map = {}
    with zipfile.ZipFile(str(SOURCE), 'r') as archive:
        for name in archive.namelist():
            if not name.startswith('word/media/'):
                continue
            target = attachments / Path(name).name
            with archive.open(name) as source, target.open('wb') as dest:
                shutil.copyfileobj(source, dest)
            image_map['/' + name] = target.name
    return image_map


def paragraph_images(paragraph):
    result = []
    for drawing in paragraph._p.xpath('.//w:drawing'):
        for blip in drawing.xpath('.//a:blip'):
            rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rid:
                part = paragraph.part.related_parts[rid]
                result.append('/' + str(part.partname).lstrip('/'))
    return result


def iter_body(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith('}p'):
            for paragraph in doc.paragraphs:
                if paragraph._p is child:
                    yield 'paragraph', paragraph
                    break
        elif child.tag.endswith('}tbl'):
            for table in doc.tables:
                if table._tbl is child:
                    yield 'table', table
                    break


def heading_level(paragraph):
    style = paragraph.style.name if paragraph.style else ''
    if style == 'Heading 1':
        return 1
    if style == 'Heading 2':
        return 2
    if style == 'Heading 3':
        return 3
    return 0


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    image_map = extract_images(Document(str(SOURCE)))
    doc = Document(str(SOURCE))
    sections = []
    current = None

    for kind, item in iter_body(doc):
        if kind == 'paragraph':
            text = clean_text(item.text)
            level = heading_level(item)
            normalized_heading = re.sub(r'\s+', ' ', text).strip()
            if level == 1 and normalized_heading in NORMALIZED_SECTION_FILES:
                filename, title = NORMALIZED_SECTION_FILES[normalized_heading]
                current = {'filename': filename, 'title': title, 'blocks': []}
                sections.append(current)
                continue
            if current is None:
                continue
            for image in paragraph_images(item):
                if image in image_map:
                    current['blocks'].append(('image', image_map[image]))
            if not text:
                continue
            if level == 2:
                current['blocks'].append(('text', '## ' + text))
            elif level == 3:
                current['blocks'].append(('text', '### ' + text))
            elif level == 1:
                current['blocks'].append(('text', '# ' + text))
            elif text.startswith('【图片占位】'):
                current['blocks'].append(('text', '> 📌 ' + text))
            else:
                current['blocks'].append(('text', text))
        else:
            if current is not None:
                table = markdown_table(item)
                if table:
                    current['blocks'].append(('table', table))

    # 删除旧的同名导出文件，避免章节边界调整后残留过期内容。
    expected = {entry[0] for entry in SECTION_FILES.values()}
    for old in OUTPUT.glob('*.md'):
        if old.name in expected:
            old.unlink()

    manifest = []
    for section in sections:
        lines = ['# ' + section['title'], '', '*从论文 Word 工作稿提取，源文件：`边坡地震动放大效应研究论文初稿（第三章重构）.docx`。*', '', '---', '']
        image_index = 0
        for block_index, (kind, value) in enumerate(section['blocks']):
            if kind == 'image':
                image_index += 1
                lines.append('![论文图形附件 {}](附件/{})'.format(image_index, value))
                lines.append('*图形附件 {}：由 Word 工作稿中的嵌入图形提取。*'.format(image_index))
            else:
                lines.append(value)
            if block_index != len(section['blocks']) - 1:
                lines.append('')
        target = OUTPUT / section['filename']
        target.write_text('\n'.join(lines).rstrip() + '\n', encoding='utf-8')
        manifest.append('- [{}]({})'.format(section['title'], section['filename']))

    index = OUTPUT / 'README.md'
    index_lines = [
        '# 论文分章节 Markdown 导出',
        '',
        '*由论文 Word 工作稿按一级标题拆分生成；图片统一保存在 `附件/`。*',
        '',
        '---',
        '',
        '## 📚 文件清单',
        '',
    ] + manifest + ['', '## 📦 附件', '', '- `附件/`：从 Word 文档中提取的嵌入图片。', '']
    index.write_text('\n'.join(index_lines), encoding='utf-8')
    print('已导出 {} 个独立 Markdown 文件，附件 {} 个。'.format(len(sections), len(image_map)))


if __name__ == '__main__':
    main()
