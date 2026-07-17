# -*- coding: utf-8 -*-
from __future__ import print_function

from copy import deepcopy
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


SOURCE = r"docs/论文材料/边坡地震动放大效应研究论文初稿.docx"
REVISED = r"docs/论文材料/边坡地震动放大效应研究论文初稿（整合修订）.docx"
ARCHIVE = r"docs/论文材料/边坡地震动放大效应研究方法验证与创新定位说明.docx"


def set_paragraph_style(paragraph, style_name):
    try:
        paragraph.style = style_name
        return
    except KeyError:
        reference = None
        for candidate in paragraph.part.document.paragraphs:
            if candidate.style and candidate.style.name == style_name:
                reference = candidate
                break
        if reference is None or reference._p.pPr is None:
            raise RuntimeError("无法复制段落样式：%s" % style_name)
        paragraph._p.insert(0, deepcopy(reference._p.pPr))


def replace_first(doc, prefix, text, style_name=None):
    matches = [p for p in doc.paragraphs if p.text.startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError("无法唯一定位段落：%s（匹配数=%d）" % (prefix, len(matches)))
    paragraph = matches[0]
    paragraph.text = text
    if style_name:
        set_paragraph_style(paragraph, style_name)
    return paragraph


def replace_nth(doc, prefix, text, index, style_name=None):
    matches = [p for p in doc.paragraphs if p.text.startswith(prefix)]
    if len(matches) <= index:
        raise RuntimeError("无法按序定位段落：%s（匹配数=%d，索引=%d）" % (prefix, len(matches), index))
    paragraph = matches[index]
    paragraph.text = text
    if style_name:
        set_paragraph_style(paragraph, style_name)
    return paragraph


def insert_after(paragraph, text, style_name=None):
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    new_paragraph = Paragraph(element, paragraph._parent)
    if style_name:
        set_paragraph_style(new_paragraph, style_name)
    new_paragraph.add_run(text)
    return new_paragraph


def add_after(paragraph, blocks):
    last = paragraph
    for text, style_name in blocks:
        last = insert_after(last, text, style_name)
    return last


def set_east_asia(run, font_name):
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    set_east_asia(run, "宋体")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_archive_paragraph(doc, text, indent=True, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.333
    paragraph.paragraph_format.space_after = Pt(8)
    if indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        set_east_asia(run, "宋体")
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_east_asia(rest, "宋体")
    else:
        run = paragraph.add_run(text)
        set_east_asia(run, "宋体")
    return paragraph


def add_archive_heading(doc, text, level):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.keep_with_next = True
    if level == 1:
        paragraph.paragraph_format.space_before = Pt(18)
        paragraph.paragraph_format.space_after = Pt(10)
        size = 14
    else:
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
        size = 12
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    set_east_asia(run, "黑体")
    return paragraph


def add_archive_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Cm(widths[index])
        set_cell_text(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "F4F6F9")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].width = Cm(widths[index])
            set_cell_text(cells[index], value)
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side in ("top", "start", "bottom", "end"):
                node = tc_mar.find(qn("w:" + side))
                if node is None:
                    node = OxmlElement("w:" + side)
                    tc_mar.append(node)
                node.set(qn("w:w"), "80" if side in ("top", "bottom") else "120")
                node.set(qn("w:type"), "dxa")
    return table


def build_revised_thesis():
    doc = Document(SOURCE)

    replace_first(doc, "坡地作为山区最普遍的不规则地形", "坡地作为山区最普遍的不规则地形，在地震作用下会显著改变地震波的传播、散射与干涉过程，产生的地形放大效应往往加重坡顶及坡缘建筑的震害。现行抗震规范对地形放大系数的规定基于均匀介质与垂直入射假定，难以涵盖实际工程中“上土下岩”成层地质结构与地震波斜入射条件下的复杂放大规律。为此，本文以斜入射SV波作用下的坡地为研究对象，采用理论解析、有限元数值模拟与机器学习相结合的方法，系统研究地形、地层与波动输入的频率—空间耦合机制。本文将线性参数化分析定位为识别该机制的零阶基准，并以等效线性强度扫描检验其在强震下的适用边界。主要工作与结论如下：")
    replace_nth(doc, "（2）建立并验证了以坡高为基准的无量纲化坡地有限元建模方法", "（2）建立并验证了以坡高为基准的无量纲化坡地有限元建模方法。模型各部位尺寸以坡高的倍数统一导出，网格按最软层波速与截止频率的Kuhlemeyer-Lysmer判据控制，人工边界弹簧—阻尼参数逐节点按所在地层取材。采用“解析/独立算法对照—文献复现—退化回归—远场一维理论—数值收敛与边界影响”五层证据链验证方法有效性：复现已发表文献的522组均质岩坡参数化工况与成层岩坡典型工况，代表性工况的时间步、侧向净空和模型深度扰动对坡顶TAF的影响分别为+0.3%、+0.1%和−1.1%，远场响应与一维理论解偏差在2%以内。", 0)
    replace_first(doc, "（3）基于522组有限元工况系统研究了均质坡地地震动放大规律", "（3）基于522组线弹性有限元工况系统研究了均质坡地地震动放大规律。结果表明：斜入射条件下水平向地形放大系数（TAF）呈显著非对称分布，最大TAF可达垂直入射工况的1.4~2.2倍；TAF最大值与坡高呈非单调关系，在坡高约为入射波长0.2倍时出现临界放大；水平向峰值加速度最大值出现在坡顶附近并向坡后平台迅速衰减，坡脚存在去放大现象。在此基础上，提出以坡高/波长比、坡角、入射角、覆盖层厚度比与波阻抗比为控制参数，并以耦合残差量化地形—地层可分离性误差的成层坡地参数化分析框架。")
    replace_first(doc, "（4）建立了基于本征正交分解", "（4）建立了基于本征正交分解（POD）与高斯过程回归（GPR）的坡地地震动放大效应代理模型。以161个地表测点的线弹性TAF曲线为目标，经POD降维后以GPR回归，按几何分组交叉验证的全区平均绝对误差为0.0244、决定系数达0.947，优于规则网格插值等非机器学习基线与CNN、LSTM等深度学习模型；通过留一外推测试诚实界定了模型适用域。该模型的贡献在于小样本全曲线代理与不确定性约束，而不以“首次机器学习预测坡地放大”作为创新表述；面向成层和强度依赖响应的扩展采用物理分解残差学习，而不将线性传递函数外推为普适的非线性解。")
    replace_first(doc, "（5）开展了地形-土-结构相互作用", "（5）开展了地形—土—结构相互作用（TSSI）的分级验证与初步分析。通过固定基础框架、平地SSI、坡顶SSI三级递进模型验证了建模与后处理流程，初步算例表明坡顶自由场放大与结构动力放大存在链式贯通效应，SSI引起的周期延长在结构自振频率接近场地卓越频率时可产生失谐减震效应。作为自由场强度依赖性的补充检验，软土等效线性扫描表明，随输入强度增大坡顶绝对TAF总体下降，但远场衰减可更快；因此，非线性并不支持“坡顶相对放大必然削弱”的一般化判断。")
    replace_first(doc, "本文研究成果可为山区坡地建筑抗震设计中地形放大系数", "本文研究成果可为山区坡地建筑抗震设计中地形放大系数的精细化取值、数值模拟结果的可信使用与快速预测提供理论依据；其工程外推仍须受材料本构、二维简化和适用域判别的共同约束。")

    replace_first(doc, "(2) A dimensionless finite element modeling procedure", "(2) A dimensionless finite element modeling procedure for slope sites scaled by the slope height is established and verified. All model dimensions are derived as multiples of the slope height, the mesh size is controlled by the Kuhlemeyer-Lysmer criterion based on the softest layer, and the spring-dashpot parameters of the artificial boundary are assigned node by node according to the local stratum. A five-level verification chain, including analytical or independent-solver checks, literature reproduction, degeneration regression, one-dimensional far-field comparison, and numerical sensitivity checks, is adopted. The procedure reproduces 522 published homogeneous-rock cases and representative layered-rock cases; the crest TAF changes by +0.3%, +0.1%, and −1.1% under the representative time-step, lateral-clearance, and model-depth perturbations, respectively, while the far-field deviation from the one-dimensional solution is within 2%.")
    replace_first(doc, "(3) The amplification behavior of homogeneous slopes", "(3) The amplification behavior of homogeneous slopes is systematically studied based on 522 linear finite element cases. The horizontal topographic amplification factor (TAF) exhibits significant asymmetry under oblique incidence, and the maximum TAF can reach 1.4–2.2 times that under vertical incidence. The maximum TAF is non-monotonic with respect to slope height and reaches a critical amplification when the slope height is about 0.2 times the incident wavelength. A layered-slope framework controlled by the height-to-wavelength ratio, slope angle, incident angle, cover-thickness ratio and impedance ratio is proposed, with a coupling residual used to quantify the error of separable topographic and stratigraphic factors.")
    replace_first(doc, "(4) A surrogate model for slope amplification", "(4) A surrogate model for slope amplification based on proper orthogonal decomposition (POD) and Gaussian process regression (GPR) is established for the linear database. Taking the TAF curves at 161 surface observation points as the target, the group-wise cross-validated mean absolute error is 0.0244 and the coefficient of determination reaches 0.947, outperforming non-machine-learning baselines and the tested deep-learning models. The applicable domain is delineated by leave-one-out extrapolation tests. The model is positioned as a small-sample full-curve surrogate with uncertainty control, rather than a claim of the first machine-learning prediction of slope amplification.")
    replace_first(doc, "(5) A step-by-step verified analysis", "(5) A step-by-step verified analysis of topography-soil-structure interaction (TSSI) is carried out. Preliminary cases show a chained amplification path from the free-field response at the slope crest to the dynamic amplification of the structure, while SSI-induced period elongation can detune the structure from site resonance. An equivalent-linear soft-soil intensity scan is further used as a boundary check: the crest response generally decreases with increasing intensity, but the far-field response may decrease faster; therefore, material nonlinearity does not imply a universal weakening of relative crest amplification.")
    replace_first(doc, "The results of this thesis provide a theoretical basis", "The results of this thesis provide a theoretically constrained basis for refined topographic-amplification assessment, credible use of numerical results and rapid prediction for buildings on mountain slopes, subject to explicit constitutive, dimensional and applicable-domain limitations.")

    replace_first(doc, "此外，坡地地震响应是一个涉及几何、地质、波动输入", "此外，坡地地震响应是一个涉及几何、地质与波动输入的高维耦合问题。这里的“耦合”首先指波场传播、散射和干涉的物理交互，不应与土体本构非线性混为一谈。传统经验拟合公式难以捕捉此类多参数映射关系；机器学习可用于构建高精度、可快速评估的代理模型，但必须以严格的数据划分、物理基线和适用域判别约束其使用。")
    replace_first(doc, "综上，地形效应与地层效应的非线性耦合", "综上，地形效应与地层效应的频率—空间耦合、斜入射引起的非对称放大以及强震下材料强度依赖性，已成为坡地地震动放大研究的前沿问题。其中，地形—地层耦合并不等同于材料本构非线性；系统识别二者可近似分离的条件及其失效区，仍缺少多参数定量研究。")
    replace_first(doc, "综上，机器学习凭借强大的非线性映射能力", "综上，机器学习在场地放大系数预测方面展现出应用潜力，但“线性有限元数据+黑箱模型预测单点放大系数”本身已不是充分创新点。现有研究多以单点放大系数或粗粒度地形参数为目标；针对坡地全地表放大曲线的代理建模、模型外推适用域的诚实评估，以及以物理分解残差刻画地形—地层耦合的数据集构建，仍有待系统推进。")
    replace_first(doc, "（1）地形效应与地层效应的非线性耦合规律", "（1）地形效应与地层效应的可分离条件及耦合残差规律尚不清楚。现行规范及多数既有研究常将地形效应与地层效应分开处理或基于均匀介质假设进行简化分析，而实际工程场地常呈现典型的“上土下岩”二元结构。覆盖层厚度、剪切波速与下伏基岩波阻抗差异会改变波的传播路径和共振频段。关键不只是判断“是否耦合”，而是量化在何种无量纲条件下可近似相乘、何时产生显著残差。")
    replace_first(doc, "（4）缺乏基于大规模数据集的高精度快速预测模型", "（4）缺乏兼具物理约束与可信度边界的快速预测模型。地形放大系数的既有预测主要依赖有限参数拟合的经验公式，难以捕捉高维交互；机器学习应用则多以单点放大系数为目标，缺乏针对全地表放大曲线的代理建模、外推适用域评估与耦合残差学习。对于强度依赖响应，若直接把后处理得到的模量折减或阻尼比作为输入，还会产生信息泄漏问题。")
    replace_first(doc, "针对上述问题，本文以斜入射SV波作用下的坡地", "针对上述问题，本文以斜入射SV波作用下的坡地（含均质岩坡与“上土下岩”成层坡地）为研究对象，采用理论解析、有限元数值模拟与机器学习相结合的方法，围绕“地层放大与地形放大在何种无量纲条件下可近似分离、何时发生显著频率—空间耦合，以及如何可信预测该耦合残差”开展研究。线弹性参数化分析承担广覆盖机制识别与基线构建的任务；等效线性强度扫描作为强震边界检验，不替代也不掩盖真非线性、孔压和失稳问题。全文共分7章，各章主要内容如下：")
    replace_first(doc, "第3章，坡地有限元模型建立与验证", "第3章，坡地有限元模型建立与验证。建立以坡高为基准的无量纲化模型尺寸设计方法；给出多层介质精确几何剖分、边界参数逐层取材、控制性最软层网格设计与频率相关瑞利阻尼层间一致标定等关键建模技术；以“文献复现—退化回归—一维理论—数值收敛—边界影响”的分层证据链验证方法有效性，并明确尚待完成的独立二维基准与一维场地软件对照。")
    replace_first(doc, "第4章，坡地地震动放大效应参数化分析", "第4章，坡地地震动放大效应参数化分析。定义地形放大系数及其归一化口径，建立以坡高/波长比、坡角、入射角、覆盖层厚度比、波阻抗比为控制参数的无量纲参数化框架与三段归一化地表坐标系统；基于522组均质工况系统分析入射角、坡高、坡角对放大效应的影响规律；以r(f,s)=ln|H_FEM|−ln|H_1D|−ln|H_topo|定义耦合残差，为成层坡地的可分离性判别和机制分区图提供统一指标。")
    replace_first(doc, "第5章，地形-土-结构相互作用分析", "第5章，地形—土—结构相互作用分析。建立“固定基础框架—平地SSI—坡顶SSI”分级验证的TSSI分析流程，采用去耦对比法定量分离SSI效应，初步揭示坡顶自由场放大向结构响应传递的链式机制；以软土等效线性强度扫描说明材料强度依赖性对该线性基线的修正方式，并给出结构真非线性与距坡缘距离影响的后续分析方案。")
    replace_first(doc, "第6章，基于机器学习的坡地地震动放大效应预测模型", "第6章，基于机器学习的坡地地震动放大效应预测模型。构建以地表TAF曲线为目标的线性代理模型数据集，建立POD降维与高斯过程回归相结合的预测框架，与物理基线、插值基线及深度学习模型对比评估精度，并通过留一外推测试界定模型适用域；进一步提出面向成层和强度依赖响应的物理分解残差学习与主动补点路线。")
    replace_first(doc, "本文的总体技术路线如图1.2所示", "本文的总体技术路线如图1.2所示：以波动理论为基础建立精确波动输入方法，以有限元数值模拟为核心生成多因素耦合的放大效应数据集；先建立线性频率—空间基准，再以等效线性扫描识别其强度依赖边界；以耦合残差而非总放大作为成层代理模型的主要学习对象，并以不确定性和适用域判别约束快速预测；最后以TSSI分析延伸考察坡顶结构的地震响应。")

    replace_first(doc, "基本假定如下：①平面应变假定", "基本假定如下：①平面应变假定——坡体沿走向长度远大于横断面尺寸，波动输入方向位于横断面内；②材料线弹性基准——本文大样本参数化研究聚焦弱至中等强度激励下的频率—空间放大机制，各层介质取线弹性本构并以瑞利阻尼计入材料衰减；③强震边界——土体强度依赖性通过独立的等效线性强度扫描补充检验，等效线性结果仅用于描述模量折减和阻尼增长的趋势，不代表真非线性、孔压累积或滑移破坏；④入射波为平面SV波，入射角不超过临界角。")

    replace_first(doc, "本文建模方法的有效性通过\"文献复现+理论对拍\"", "本文建模方法的有效性不以单一文献曲线吻合为依据，而采用“解析/独立算法对照—文献复现—模型退化—远场理论—数值收敛与边界影响”的分层验证方案。每一层分别约束波动输入正确性、复杂坡地响应的外部一致性、脚本迭代不引入偏差、观测窗的理论退化性以及离散与截断误差。对尚缺少的独立二维基准和一维场地软件交叉对照，本文明确列为结论外的待闭环工作。")
    sensitivity_anchor = replace_first(doc, "模型尺寸设计（第3.3节）的有效性由逐工况", "模型尺寸设计（第3.3节）的有效性由逐工况的自动化定量校核保证：远离坡体的模型左、右边界附近，地表响应应退化为一维水平成层自由场。本文在每个工况建模时，同步以频域传递矩阵法计算左、右边界土柱的一维粘弹性自由场理论解并写入工况元数据；后处理时将有限元远场台阶值与理论值对比，偏差超过5%的工况自动告警并排查。既有算例中该偏差均在2%以内，表明本文尺寸设计对观测区精度是充分的，边界净空、阻尼标定与网格设计均满足要求。")
    add_after(sensitivity_anchor, [
        ("3.8.5 数值收敛与边界影响检验", "Heading 3"),
        ("在代表性成层坡地工况上，分别实施时间步减半、侧向边界外移和底部加深的单因素扰动，以坡顶TAF为统一判据。相对基准模型，时间步减半后坡顶TAF变化为+0.3%，侧向边界外移后变化为+0.1%，底部加深后变化为−1.1%，均小于预设的5%工程容许阈值。该结果说明当前网格—时间步—计算域组合对目标观测区已达到数值收敛；这里的结论仅针对已检验的频带、材料与几何范围，不能外推为所有极端工况的通用证明。", "Body Text CN"),
        ("对粘弹性人工边界另行开展反射性检验。结果显示，源点处剩余反射约为2.6%；在距源+120 m和+180 m的目标观测区，剩余反射分别约为15.4%和12.1%，显著低于固定边界对应的67.3%和83.3%。因此，本文不将该边界表述为“完全无反射”，而以目标观测区对边界处理不敏感、并通过远场一维理论对拍作为可接受性判据。", "Body Text CN"),
        ("3.8.6 验证边界与后续闭环", "Heading 3"),
        ("现有证据已覆盖自由场算法、文献工况、退化回归、远场理论、数值收敛及边界影响，足以支撑本文线性参数化研究的内部可信性。为进一步增强外部有效性，后续应选取一组不参与参数标定的二维地形公开基准进行盲预测，并以DEEPSOIL等一维场地响应软件交叉校核平坦成层场地；这两项属于论文提交前应优先补强的独立验证，而非当前已完成的结果。", "Body Text CN")
    ])
    replace_first(doc, "（4）通过复现Shen等[7]的522组均质岩坡", "（4）采用分层证据链验证了模型可信性：522组均质岩坡与成层岩坡典型工况复现用于外部规律对照，均质退化回归与远场一维理论对拍（偏差<2%）用于内部自洽性检验，代表性工况的时间步、侧向净空和模型深度扰动分别仅引起+0.3%、+0.1%和−1.1%的坡顶TAF变化；粘弹性边界在目标观测区的剩余反射显著低于固定边界。独立二维公开基准和DEEPSOIL交叉对照仍列为后续闭环任务。")

    residual_anchor = replace_first(doc, "成层坡地的放大是地形效应与地层效应的耦合", "成层坡地的放大是地形效应与地层效应的波场耦合：地形效应主要由η、i、θ控制，地层效应主要由d/h、α_z控制，二者通过散射、界面透反射和多次波干涉相互影响，不能预设为简单相乘[66,67]。为避免将“耦合”笼统等同于材料本构非线性，本文在线性频域框架内先回答其可分离性问题：以H_FEM(f,s)为成层坡地完整传递函数、H_1D(f)为同一成层平地的一维传递函数、H_topo(f,s)为相应均质地形传递函数，比较三者的偏离。覆盖层基频f0=V_s1/(4d)与坡体地形特征频率接近时可能出现双重共振；阻抗比α_z和厚度比d/h则控制共振峰及其与地形特征频带的相对位置。")
    add_after(residual_anchor, [
        ("耦合残差定义为 r(f,s)=ln|H_FEM(f,s)|−ln|H_1D(f)|−ln|H_topo(f,s)|。当r接近零时，地层与地形放大可在该频率—空间位置近似分离；当|r|显著增大时，简单相乘的近似失效。后续以η、i、θ、d/h和α_z为坐标建立“近似可分离—强耦合”的机制分区图，并以未参与拟合的工况检验其判别能力。", "Body Text CN")
    ])
    replace_first(doc, "首批78组成层工况", "首批78组成层工况（坡角×覆盖层厚度×波阻抗比交叉组合）已完成建模方案设计与脚本验证（第3.8.2节），批量计算仍在推进。本节将已完成的均质线性结果作为核心结论；成层“可分离—强耦合”分区图必须待独立留出工况验证后再写为正式规律。同步开展的软土等效线性强度扫描仅用于检验线性基准的强震边界，不能与本节的线性耦合残差混为同一数据池。")
    replace_first(doc, "【图片占位】图4.9占位", "【图片占位】图4.9占位——地形×地层耦合残差图：r(f,s)=ln|H_FEM|−ln|H_1D|−ln|H_topo|在频率—空间—参数域的分布，量化“解耦相乘假定”的误差")
    replace_first(doc, "其中图4.9对应的耦合残差分析具有直接的规范意义", "图4.9对应的耦合残差分析首先服务于机制判别：残差显著偏离零的参数区即“地形系数与场地系数简单相乘”近似可能失效的区域。只有在完成成层工况、独立验证和工程参数映射后，才可进一步讨论其对规范取值的启示；本研究不以尚未完成的数据直接提出规范修订值。")
    replace_first(doc, "（4）初步规范对比表明", "（4）初步规范对比表明，斜入射叠加临界高度的不利组合下数值放大系数可超出GB 50011上限约25%以上，提示线性均质工况下的规范包络范围需要进一步核验；成层坡地的双重共振、耦合残差分区和强震强度依赖修正仍须待批量计算与独立验证完成后再形成工程性结论。")

    nonlinear_head = replace_first(doc, "5.5 后续分析方案：结构非线性与距坡缘距离效应", "5.5 强度依赖性及结构非线性的研究边界")
    replace_first(doc, "本章已验证的弹性TSSI流程之上，规划两步深化分析", "在本章已验证的弹性TSSI流程之上，需把“线性主研究”“等效线性边界检验”和“结构真非线性扩展”清晰分层，避免以一种层级的结论替代另一种层级。")
    nonlinear_structure = replace_first(doc, "（1）结构非线性。将梁柱替换", "5.5.2 结构非线性", "Heading 3")
    add_after(nonlinear_structure, [
        ("将梁柱替换为纤维截面或集中塑性铰模型（混凝土塑性损伤本构），考察强震下结构刚度退化与塑性损伤对TSSI响应的影响，输出损伤指标的空间分布。技术风险在于隐式积分与材料非线性的收敛性，必要时转为显式积分并重新校核斜入射等效节点力的施加方式。需要说明的是，坡地自由场采用频域方法（第2章）与结构真非线性的组合在“土体等效线性+结构非线性”的框架下可以自洽；若土体亦需真非线性，则波动输入需改用时域自由场算法，此为方法边界。", "Body Text CN")
    ])
    nonlinear_distance = replace_first(doc, "（2）距坡缘距离扫描。以框架距坡肩距离", "5.5.3 距坡缘距离效应", "Heading 3")
    add_after(nonlinear_distance, [
        ("以框架距坡肩距离与结构基础宽度之比M/T为参数开展系列工况，量化坡顶后缘不同位置结构的响应差异与破坏模式演化，结合第4章TAF空间分布曲线，给出坡顶建筑退让距离的量化建议。", "Body Text CN")
    ])
    add_after(nonlinear_head, [
        ("5.5.1 等效线性强度扫描及其解释", "Heading 3"),
        ("为判断线性主研究在强震下的边界，对代表性软土坡地实施等效线性（EQL）强度扫描：以应变相容的模量折减和阻尼增长迭代更新介质参数，考察0.05g、0.10g、0.20g、0.40g和0.80g输入下的峰值响应。线性基准的坡顶与远场TAF分别为2.420和2.051；扫描中坡顶TAF依次为2.429、2.374、2.211、1.906和1.692，远场TAF依次为1.929、1.795、1.549、1.277和1.055。", "Body Text CN"),
        ("该代表性结果显示，随强度升高，坡顶绝对响应总体降低，体现了剪切模量下降、阻尼增大及共振频带迁移的共同作用；但坡顶/远场的相对比值由线性基准约1.18提高至0.80g时约1.60。这说明远场一维场地响应在该模型中衰减更快，不能据“非线性使坡顶绝对响应下降”直接推出“坡顶相对地形放大必然削弱”。不同土层、频带和输入谱形下的趋势仍需统计检验。", "Body Text CN"),
        ("EQL适用于本研究中强度依赖放大趋势的对比，不等价于真非线性动力分析：它不能可靠描述循环累积塑性、孔隙水压力发展、接触开裂、滑移和失稳。故本文仍以线性频域H(f,s)作为广覆盖机制和机器学习基线；当研究目标转向破坏、孔压或永久位移时，应另建真非线性时域分析链。", "Body Text CN")
    ])
    replace_first(doc, "（3）坡顶初步算例揭示了", "（3）坡顶初步算例揭示了“地形地层放大（3.89倍）→结构动力放大（顶层4.67倍）”的链式机制，坡顶结构响应较平地对照高出近3倍，与震害观测定性吻合。代表性软土EQL扫描进一步表明：强震下坡顶绝对响应可降低，但相对于远场的放大并不必然减弱；因此，线性TSSI结论应在其频带、强度和本构边界内使用，结构非线性与距坡缘距离效应仍作为后续工作实施。")

    replace_first(doc, "第4章的参数化分析揭示了坡地放大效应的主控规律", "第4章的参数化分析揭示了坡地放大效应的主控规律，但有限元逐工况计算耗时以小时计，无法满足工程快速评估与大范围风险普查的需求；传统经验公式又难以捕捉多参数的高维耦合。本章以线弹性有限元数据集为基础，构建“降维+回归”的代理模型：以地表161个测点的完整TAF曲线（而非单点放大系数）为预测目标，采用本征正交分解（POD）压缩目标维度，以高斯过程回归（GPR）与XGBoost建立特征到POD系数的映射；通过与物理、插值及深度学习基线的系统对比评估精度；并以留一外推测试诚实界定模型的适用域。第6.6节再讨论其向成层与强度依赖响应扩展的必要条件。")
    replace_first(doc, "本章数据集来自第3、4章的均质坡地有限元工况", "本章数据集来自第3、4章的均质坡地线弹性有限元工况：174组几何工况（坡角×坡高等交叉组合）×3条实际地震波（El Centro、Loma Prieta、Northridge，PGA统一调幅）＝522组样本。每组样本的目标为沿坡面地表161个空间测点的TAF曲线，特征为10维输入向量。该数据集用于建立线性零阶代理，而不是宣称已经覆盖强震土体真非线性。")
    replace_first(doc, "10维输入特征分为两组", "当前10维输入特征分为两组：几何特征（坡角、坡高、坡脚距、坡顶宽4维）与地震波特征（PGA、卓越周期，及T=0.1 s、0.5 s、1.0 s处的反应谱值等6维）。地震波谱特征使模型能够区分不同频谱成分输入下同一几何的响应差异，是从“逐波建模”走向“跨波泛化”的必要尝试；但仅有3条波形尚不足以证明任意波形外推。面向成层与强度依赖扩展时，应以η、d/h、α_z及PGA、Arias强度、显著持时、Sa(T_site)等可在计算前获得的变量补充特征，而不得直接把后处理得到的G/G_max、阻尼比或有效应变作为输入。")
    replace_first(doc, "为客观定位机器学习的增益", "为客观定位机器学习的增益，设置三级非机器学习基线：B1全局均值基线（所有训练曲线取平均，代表“盲猜”下限）；B2 Oracle几何均值基线（已知测试几何时取该几何训练波的均值，代表不建模输入波影响的上限）；B3规则网格插值基线（在几何参数网格上线性插值，代表传统查表法）。同时与前期开展的时序深度学习方案（CNN、LSTM、Transformer、DeepONet，以时程为输入）对比。该比较旨在说明小样本下的模型选择，而不作为“机器学习首次用于坡地放大预测”的论据。")
    replace_first(doc, "各模型在GroupKFold下的精度对比见表6.2", "各模型在GroupKFold下的精度对比见表6.2。GPR以全区平均绝对误差0.0244、决定系数0.9474居首，XGBoost次之；二者均显著优于全部非机器学习基线，其中相对B3规则网格插值（工程查表法的代表），GPR误差降低约66%，说明其能够有效拟合当前线性数据中的参数交互。该性能结果构成应用层贡献；论文的核心科学问题仍是地形—地层耦合残差的机制及其可预测性。")
    replace_first(doc, "6.6 面向成层坡地的模型升级方案", "6.6 面向成层与强度依赖坡地的模型升级方案")
    replace_first(doc, "针对第6.5节暴露的外推短板与成层坡地的扩展需求", "针对第6.5节暴露的外推短板与成层、强度依赖坡地的扩展需求，规划四项升级。其实施顺序应遵循“先完成成层线性耦合残差闭环，再增加强度维度”的原则，避免因数据量不足而把材料非线性误判为几何或谱形效应。")
    replace_first(doc, "（2）频域目标与二维POD", "（2）频域目标与二维POD。以三段归一化坐标上的传递函数曲面H(f,s)替代时域TAF曲线作为主目标，对“空间×频率”矩阵作二维POD压缩（预计20~40阶模态）；峰值放大系数AR_max作为标量目标单独精确回归，避免重采样削峰。该H(f,s)与输入波频谱卷积的重构仅适用于线性或经单次等效线性更新后被冻结的体系；对强震全过程真非线性响应，不能将线性传递函数视为普适算子。")
    residual_ml = replace_first(doc, "（3）物理分解残差学习", "（3）物理分解残差学习。不让模型直接学习总放大，而是学习耦合残差r(f,s)=ln|H_FEM(f,s)|−ln|H_1D(f)|−ln|H_topo(f,s)|：其中一维层状理论解H_1D毫秒级可算，均质地形解H_topo由本章模型提供，残差项量级小、变化相对平缓，小样本更易学习。残差显著偏离零的参数区即“规范解耦相乘假定”可能失效区，与第4.4节的机理分析互为印证。")
    add_after(residual_ml, [
        ("（4）强度依赖残差与主动补点。在线性残差模型形成稳定基线后，以PGA、Arias强度、显著持时、Sa(T_site)及可在计算前获得的土层参数作为输入，学习非线性响应相对线性基线的修正量；若需要使用G/G_max、阻尼比或有效应变，应将其放入由输入强度预测的第一阶段模型，不能直接作为第二阶段预测器的特征。建议至少形成40~60组覆盖强度—地层—几何交互的EQL对比工况，并按模型不确定性与耦合残差优先级主动补点；在达到该规模前，本文只将P3e强度扫描用于物理机理对比，不报告“非线性机器学习”精度。", "Body Text CN")
    ])
    replace_first(doc, "【图片占位】图6.4占位", "【图片占位】图6.4占位——升级方案框架图：无量纲特征+H(f,s)二维POD+耦合残差学习+强度依赖修正的两阶段数据流（待成层数据与EQL对比样本充实后补充精度）")
    replace_first(doc, "（4）确立了面向成层坡地的升级路线", "（4）确立了面向成层与强度依赖坡地的升级路线：无量纲特征、频域传递函数曲面二维POD与物理分解残差学习分别针对性化解坡高外推、波形外推与小样本耦合学习短板；强度依赖扩展采用“线性基线+非线性修正”的两阶段思路，并以40~60组以上EQL对比工况及主动补点作为建模前提。")

    replace_first(doc, "本文以斜入射SV波作用下的坡地为研究对象", "本文以斜入射SV波作用下的坡地为研究对象，围绕地形、地层与波动输入多因素的频率—空间耦合，开展了波动输入方法、有限元建模、线性参数化分析、强度依赖性边界检验、机器学习预测与地形—土—结构相互作用五个层面的研究。本文将线性模型作为广覆盖的零阶基准，将等效线性扫描作为强震下的对照，不以未完成的真非线性计算替代已验证的线性结论。主要工作与结论如下：")
    replace_nth(doc, "（2）建立并验证了以坡高为基准的无量纲化坡地有限元建模方法", "（2）建立并验证了以坡高为基准的无量纲化坡地有限元建模方法。模型各部位尺寸以坡高倍数统一导出、几何与地层配置解耦，网格按控制性最软层与Kuhlemeyer-Lysmer判据控制，人工边界参数逐节点取材，瑞利阻尼逐层双频点标定。通过文献复现、均质退化、远场一维理论、数值收敛和边界影响构成的证据链完成验证；代表性工况的时间步减半、侧向外移与底部加深仅使坡顶TAF变化+0.3%、+0.1%和−1.1%，远场响应与一维理论解偏差在2%以内。", 1)
    replace_first(doc, "（4）建立了以地表TAF全曲线为目标", "（4）建立了以地表TAF全曲线为目标的线性机器学习代理模型。POD降维（15阶，累计方差99.5%）与高斯过程回归相结合，在按几何分组的严格交叉验证下达到全区MAE 0.0244、R² 0.947，显著优于规则网格插值等基线与测试的深度学习模型，训练成本CPU秒级、预测毫秒级。留一外推测试诚实界定了适用域。该结果说明在当前小样本线性数据上“降维+经典回归”优于端到端深度学习，但其创新定位应与耦合残差机制和可信度约束结合，而非止于直接预测。")
    innovation_second = replace_nth(doc, "（2）系统剖析了时域射线叠加类波动输入方法", "（2）系统剖析了时域射线叠加类波动输入方法在多层软覆盖场地下的六类误差来源，证明其中口径错配类误差因双程乘积对称性在常规单层校核中不可见，并定量给出其对软表层场地边界输入幅值的低估幅度，为同类方法的适用性判断提供了机理依据。", 1)
    add_after(innovation_second, [
        ("（3）提出以耦合残差r(f,s)=ln|H_FEM|−ln|H_1D|−ln|H_topo|为核心的成层坡地研究范式。该范式将科学问题由“某参数是否放大”推进为“地层放大与地形放大何时可近似分离、何时在特定频率—空间位置显著耦合”，并以[η,i,θ,d/h,α_z]无量纲坐标、三段归一化地表坐标和独立留出工况共同支撑机制分区图。当前论文已完成残差定义、计算管线和线性基准，成层分区结论须待批量计算与外部验证闭环后固化。", "Body Text CN")
    ])
    replace_nth(doc, "（3）建立了\"无量纲参数框架+三段归一化地表坐标+POD降维", "（4）建立了“无量纲参数框架+三段归一化地表坐标+POD降维+高斯过程回归”的坡地放大效应代理建模方法，以全地表放大曲线（而非单点系数）为预测目标，并以按几何分组交叉验证、留一外推测试和不确定性预警诚实界定模型适用域；进一步提出物理分解残差学习和强度依赖修正路线。该贡献定位为可解释、可拒识的小样本工程代理建模，而非宣称机器学习预测地形放大的首次应用。", 0)
    replace_first(doc, "需要说明的是，本文波动输入方法在算法层面继承", "需要说明的是，本文波动输入方法在算法层面继承了Thomson-Haskell/全局矩阵的经典框架，创新点（1）（2）的贡献在于面向粘弹性边界等效输入的工程实现、衰减一致化处理与误差机理的系统揭示，而非频域精确解本身；创新点（3）以可检验的耦合残差与机制分区为核心，必须以留出工况和独立基准支撑，不能以概念提出替代证据。")
    replace_first(doc, "（1）成层坡地参数化数据尚未批量完成", "（1）成层坡地参数化数据尚未批量完成。78组多层工况的批量计算正在进行，第4.4节的可分离—强耦合机制分区、耦合残差量化与规范对比有待数据完成后系统展开；后续按主动学习策略补点至200~300组，并保留独立留出工况检验分区图和残差代理模型。")
    replace_first(doc, "（2）材料本构与维度简化", "（2）材料本构与维度简化。本文以线弹性+瑞利阻尼构成参数化主研究，以代表性软土EQL强度扫描检验模量折减和阻尼增长下的边界。当前扫描显示坡顶绝对TAF随强度增大总体下降，而坡顶/远场相对比可增加，故不存在“非线性必然削弱坡顶放大”的一般结论；EQL亦不能取代循环塑性、孔压、接触开裂和滑移失稳的真非线性分析。二维平面应变模型的三维效应、坡体走向与入射方位仍有待研究。")
    replace_first(doc, "（4）真实地震案例检验不足", "（4）独立外部验证与真实地震案例检验不足。除进一步选取2~3个有详实震害调查与场地资料的案例外，论文提交前应优先补充一组不参与参数标定的二维地形公开基准和DEEPSOIL等一维场地软件对照，用以区分“内部数值自洽”与“外部方法有效”。")

    doc.save(REVISED)


def build_archive():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.space_after = Pt(8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("边坡地震动放大效应研究\n方法有效性、非线性与创新定位说明")
    run.bold = True
    run.font.size = Pt(18)
    set_east_asia(run, "黑体")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("供论文修订、开题/中期检查与答辩沟通使用")
    run.font.size = Pt(10.5)
    set_east_asia(run, "宋体")

    add_archive_heading(doc, "1 归档结论", 1)
    add_archive_paragraph(doc, "本研究的切入点仍然成立，但需要从“线性有限元+机器学习直接预测TAF”调整为“斜入射成层坡地中地层放大与地形放大的频率—空间耦合机制、可分离条件及其物理约束预测”。线性有限元不是对真实强震的替代，而是用于构建机制清晰、可大范围扫描、可形成频域基线的零阶模型；等效线性强度扫描用于检验该基线何时开始受到材料强度依赖性的显著修正。")
    add_archive_paragraph(doc, "论文不应宣称“首次用机器学习预测坡地放大”，也不应将粘弹性人工边界描述为“完全无反射”。更稳健且可审查的贡献是：建立可复核的数值验证链；以耦合残差量化地形与地层相乘近似的失效；以POD—GPR作为小样本全曲线代理，并附带不确定性和外推预警。", bold_prefix="论文不应")

    add_archive_heading(doc, "2 有限元方法有效性：应写成证据链，而非一句“与文献吻合”", 1)
    add_archive_paragraph(doc, "数值方法有效性应同时回答三个问题：波动输入是否正确、截断与离散误差是否受控、复杂坡地结果是否与独立知识一致。下表将现有证据与尚待完成的外部闭环分开列示，论文中只能将“已完成”项写成结论。")
    add_archive_table(doc,
        ["层级", "核验对象", "现有证据/定量结果", "状态与写作口径"],
        [
            ["I", "自由场与波动输入", "半空间解析退化；成层垂直入射与独立Thomson—Haskell对照。", "已完成；证明输入算法正确性。"],
            ["II", "外部坡地响应", "复现522组均质岩坡参数化规律及成层岩坡典型工况。", "已完成；写“规律和曲线一致”，不写“完全证明”。"],
            ["III", "脚本退化与远场", "多层脚本退化回均质基线；远场与一维理论偏差<2%。", "已完成；证明内部自洽与观测窗合理。"],
            ["IV", "收敛与截断", "代表性工况：时间步减半+0.3%，侧向外移+0.1%，底部加深−1.1%。", "已完成；只对已检验频带和参数范围负责。"],
            ["V", "人工边界", "VAB剩余反射：源点2.6%；目标区+120 m为15.4%、+180 m为12.1%，固定边界对应67.3%、83.3%。", "已完成；表述为“目标区可接受”，不称“零反射”。"],
            ["VI", "独立外部闭环", "二维公开地形基准盲预测；平坦成层场地与DEEPSOIL等软件交叉对照。", "待完成；应列为提交前优先补强项。"]
        ],
        [0.75, 2.0, 7.0, 4.2])
    add_archive_paragraph(doc, "审查时的推荐逻辑是：先用层级I证明输入正确，再用层级III—V说明模型对数值设置不敏感，最后以层级II和VI说明结果具有外部可迁移性。任何一层缺失都不必否定整体研究，但必须在论文中准确标注其证据等级和适用范围。")

    add_archive_heading(doc, "3 非线性问题：它修正结论边界，不会否定研究问题", 1)
    add_archive_paragraph(doc, "代表性软土EQL扫描的线性基准坡顶/远场TAF为2.420/2.051。随输入PGA从0.05g增至0.80g，坡顶TAF由2.429降至1.692，远场TAF由1.929降至1.055；相对坡顶放大（坡顶/远场）却由线性基准约1.18升至约1.60。该组结果的正确解读不是“非线性使坡顶放大消失”，而是：在此土层、谱形和频带下，模量折减、阻尼增长与共振迁移使远场一维场地响应衰减得更快。")
    add_archive_table(doc,
        ["输入PGA", "坡顶TAF", "远场TAF", "坡顶/远场"],
        [["线性基准", "2.420", "2.051", "1.18"], ["0.05g", "2.429", "1.929", "1.26"], ["0.10g", "2.374", "1.795", "1.32"], ["0.20g", "2.211", "1.549", "1.43"], ["0.40g", "1.906", "1.277", "1.49"], ["0.80g", "1.692", "1.055", "1.60"]],
        [2.3, 3.2, 3.2, 3.2])
    add_archive_paragraph(doc, "EQL可用于强度依赖的趋势与对比，但不能可靠模拟循环累积塑性、孔隙水压力、接触开裂、滑移和失稳。因而，本文采用“线性主研究+EQL边界检验”的层级是合理的：前者服务于频率—空间机制和大样本数据，后者避免把线性结果误写为强震通用结论；若研究目标转为永久位移或失稳，则必须使用真非线性时域链。")

    add_archive_heading(doc, "4 创新定位：从“预测一个系数”转向“解释并预测耦合残差”", 1)
    add_archive_paragraph(doc, "核心科学问题可表述为：在斜入射成层坡地中，地层放大与地形放大在何种无量纲条件下可近似分离，何时发生显著耦合，以及该耦合如何沿频率—空间传播。建议使用耦合残差 r(f,s)=ln|H_FEM(f,s)|−ln|H_1D(f)|−ln|H_topo(f,s)|，其中H_FEM为完整成层坡地传递函数，H_1D为同一成层平地的一维传递函数，H_topo为相应均质地形传递函数。r≈0对应近似可分离，|r|显著对应简单相乘近似失效。")
    add_archive_table(doc,
        ["层级", "可写的创新", "必须避免的表述"],
        [
            ["核心科学", "以[η,i,θ,d/h,α_z]及三段归一化坐标构建“近似可分离—强耦合”机制分区，并用留出工况检验。", "“首次发现地形与地层耦合”。"],
            ["方法证据", "频域全局矩阵自由场、衰减一致化、分层V&V和全表面H(f,s)数据管线。", "“提出了频域全局矩阵法”。"],
            ["应用代理", "POD—GPR全曲线残差代理、预测方差/OOD预警、主动学习补点。", "“首次用机器学习预测坡地放大”。"]
        ],
        [1.8, 7.6, 4.5])
    add_archive_paragraph(doc, "论文的比较设计应覆盖：直接几何—频率回归（代表已有经验式/直接ML）、简单H_1D×H_topo相乘、黑箱代理和物理残差代理。只有在同一留出集上证明残差模型同时改善精度、可解释性和适用域提示，机器学习部分才会成为核心问题的有效支撑，而非独立卖点。")

    add_archive_heading(doc, "5 机器学习路线：线性基线与非线性修正分开建模", 1)
    add_archive_paragraph(doc, "当前522组线性数据适合POD—GPR：161点TAF曲线经15阶POD压缩，按几何分组交叉验证得到MAE 0.0244、R² 0.947。该结果说明小样本、昂贵有限元数据下，经典模型优于测试的端到端深度模型；它不是非线性土体响应模型。")
    add_archive_paragraph(doc, "向成层与强度依赖扩展时，第一阶段以η、i、θ、d/h、α_z及输入PGA、Arias强度、显著持时、Sa(T_site)等事前可得变量预测线性基线或EQL修正；第二阶段可学习耦合残差或强度修正。G/G_max、阻尼比和有效应变若来自计算后处理，不能直接作为输入，否则会构成信息泄漏；如确有必要，应先建立其预测子模型。")
    add_archive_paragraph(doc, "建议在现有P3e强度扫描之外，至少形成40~60组覆盖强度—地层—几何交互的EQL对比样本，并按预测方差和耦合残差优先级主动补点。在达到该样本量和独立测试要求前，只把非线性结果用于物理对比，不报告“非线性机器学习”的泛化精度。")

    add_archive_heading(doc, "6 本次论文修订已落实的要点", 1)
    add_archive_table(doc,
        ["论文位置", "已落实的修订"],
        [
            ["摘要与第1章", "明确线性主研究、EQL边界检验和核心科学问题；删除“ML首次应用”式隐含表述。"],
            ["第3章", "补入五层有效性证据链、收敛定量结果、VAB残余反射口径及独立验证待办。"],
            ["第4章", "把“耦合”与“本构非线性”区分；用r(f,s)定义可检验的耦合残差。"],
            ["第5章", "补入软土EQL强度扫描及其正确解释；明确EQL与真非线性的边界。"],
            ["第6章", "将代理模型限定为线性基线；新增物理残差、两阶段强度修正和主动补点条件。"],
            ["第7章", "把创新点重排为机制—方法证据—可信代理，并将未完成验证明确列为展望。"]
        ],
        [3.0, 11.0])

    add_archive_heading(doc, "参考文献定位", 1)
    add_archive_paragraph(doc, "[1] RIZZITANO S, CASCONE E, BIONDI G. Coupling of topographic and stratigraphic effects on seismic response of slopes through 2D linear and equivalent linear analyses[J]. Soil Dynamics and Earthquake Engineering, 2014.（论文原参考文献[66]）", indent=False)
    add_archive_paragraph(doc, "[2] SHEN H, LIU Y, LI X, et al. The combined amplification effects of topography and stratigraphy of layered rock slopes under vertically and obliquely incident seismic waves[J]. Soil Dynamics and Earthquake Engineering.（论文原参考文献[65]）", indent=False)
    add_archive_paragraph(doc, "[3] 罗麒锐, 赵仕兴, 吴启红, 等. 基于地形参数的地震动放大系数数值分析与预测方法[J]. 建筑结构, 2025, 55(2): 50-57, 35.（论文原参考文献[87]）", indent=False)
    add_archive_paragraph(doc, "注：本说明归档的是本轮论文定位与方法论口径。具体数值应以相应自动化验证报告和可复现算例为准；未完成项不得在论文中表述为既有结论。", indent=False)
    doc.save(ARCHIVE)


if __name__ == "__main__":
    build_revised_thesis()
    build_archive()
    print("已生成：%s" % REVISED)
    print("已生成：%s" % ARCHIVE)
