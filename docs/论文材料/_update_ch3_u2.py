# -*- coding: utf-8 -*-
"""将通过 U2 验证的输入与人工边界内容写入论文修订稿第3.4节。"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm


DOCX_PATH = r"C:\Users\12462\Documents\Code\AbqScripts\docs\论文材料\边坡地震动放大效应研究论文初稿（整合修订）.docx"


def insert_paragraph_after(document, anchor, text, style=None, alignment=None):
    """在指定 XML 元素后插入段落并返回新段落。"""
    paragraph = OxmlElement("w:p")
    anchor.addnext(paragraph)
    from docx.text.paragraph import Paragraph
    result = Paragraph(paragraph, document._body)
    if style:
        result.style = style
    result.add_run(text)
    if alignment is not None:
        result.alignment = alignment
    return result


def set_cant_split(row):
    """禁止表格行跨页拆分。"""
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_repeat_header(row):
    """把首行标记为跨页重复表头。"""
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def set_table_borders(table):
    """为验证结果表添加统一细实线边框。"""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement("w:" + edge)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tbl_pr.append(borders)


def main():
    document = Document(DOCX_PATH)
    start = next(p for p in document.paragraphs if p.text.strip().startswith("3.4 "))
    end = next(p for p in document.paragraphs if p.text.strip().startswith("3.5 "))
    start.text = "3.4 波动输入与人工边界"

    node = start._p.getnext()
    while node is not None and node is not end._p:
        following = node.getnext()
        node.getparent().remove(node)
        node = following

    paragraphs = [
        "本研究采用第2章建立的频域自由场引擎生成斜入射 SV 波场。对每个人工边界节点，程序依据其坐标、所在材料和水平慢度计算自由场位移、速度与应力时程，并将粘弹性边界反力与自由场面力组合为等效节点荷载。左、右边界的法向自由度为水平方向、切向自由度为竖直方向，底边界则相反；角点同时接受相邻两条边界的贡献。边界弹簧、阻尼器和等效荷载均由同一材料参数与节点影响长度计算，避免输入波场与吸收边界采用不同介质口径。",
        "F_eq(t)=K_b·u_ff(t)+C_b·v_ff(t)+A_b·σ_ff(t)·n",
        "式中，K_b 和 C_b 分别为节点影响范围内的边界弹簧刚度与阻尼系数，A_b 为二维模型中的节点影响长度，u_ff、v_ff 和 σ_ff 为自由场位移、速度与应力，n 为外法向。频域自由场在建模前自动执行均质半空间和单层场地解析对拍，本次两项相对误差分别为1.11×10⁻¹⁶和3.55×10⁻¹⁶。",
        "斜入射条件下，不能同时把坡体两侧端点的全时程峰值都视为一维自由场。入射波与坡体相互作用后，下游平台包含真实的地形散射波；若仍要求其 TAF_h 必须趋于1，会把物理散射误判为边界输入误差。为此，本研究采用双向斜入射配对验证：+15° 工况以左侧上游端检验，−15° 工况以右侧上游端检验；下游端误差继续记录，用于描述散射影响，但不参与输入边界是否正确的判定。垂直入射时没有水平传播方向，两端仍须同时通过。",
        "配对工况均采用均质基岩、坡高25 m、坡角30°、侧向净空6h、坡脚以下深度5h和4 Hz Ricker 波，仅改变入射角符号。+15° 工况的左上游误差为0.79%，−15° 工况的右上游误差为1.82%，均小于预设2%门槛。相应下游误差为7.07%和5.40%，且随传播方向交换左右，证明该偏差不是固定侧边界的符号或系数错误，而是与坡体散射及全时程峰值统计有关。由此，左右人工边界在分别作为传播上游端时均通过一致性验证。",
        "上述验证仅说明当前线弹性、小应变、二维平面应变模型中的波动输入与人工边界实现满足本章参数化研究要求。下游端偏差不得作为边界误差删去，也不得用上游检验结果宣称坡地远场处处等同于一维自由场；它将在后续坡地响应分析中作为传播方向效应保留。计算域、网格、时间步与阻尼对目标响应的影响仍需由第3.5节的独立收敛试验确定。"
    ]

    anchor = start._p
    for index, text in enumerate(paragraphs):
        alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 1 else WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph = insert_paragraph_after(document, anchor, text, "Body Text CN", alignment)
        anchor = paragraph._p

    caption = insert_paragraph_after(document, anchor, "表3.2 斜入射输入与人工边界双向配对验证结果", "Figure Caption CN", WD_ALIGN_PARAGRAPH.CENTER)
    anchor = caption._p
    table = document.add_table(rows=3, cols=6)
    set_table_borders(table)
    headers = ["入射角", "传播上游端", "上游误差", "传播下游端", "下游误差", "判定"]
    values = [
        ["+15°", "左端", "0.79%", "右端", "7.07%", "通过"],
        ["−15°", "右端", "1.82%", "左端", "5.40%", "通过"]
    ]
    for column, text in enumerate(headers):
        table.cell(0, column).text = text
    for row_index, row_values in enumerate(values, 1):
        for column, text in enumerate(row_values):
            table.cell(row_index, column).text = text
    for row in table.rows:
        set_cant_split(row)
    set_repeat_header(table.rows[0])
    widths = [Cm(1.6), Cm(2.4), Cm(2.0), Cm(2.4), Cm(2.0), Cm(1.6)]
    for row in table.rows:
        for column, width in enumerate(widths):
            row.cells[column].width = width
    anchor.addnext(table._tbl)

    document.save(DOCX_PATH)
    print("已写入第3.4节并保存：%s" % DOCX_PATH)


if __name__ == "__main__":
    main()
