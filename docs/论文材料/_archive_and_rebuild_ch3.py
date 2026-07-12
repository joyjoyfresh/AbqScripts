# -*- coding: utf-8 -*-
"""归档重构前第三章，并生成不含旧验证结论的第三章工作稿。"""

from __future__ import print_function

import hashlib
import os
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
SOURCE = os.path.join(ROOT, 'docs', '论文材料', '边坡地震动放大效应研究论文初稿（整合修订）.docx')
ARCHIVE_CH3 = os.path.join(ROOT, 'docs', '论文材料', '归档', '2026-07-12_第三章重构前正文.docx')
WORKING = os.path.join(ROOT, 'docs', '论文材料', '边坡地震动放大效应研究论文初稿（第三章重构）.docx')

OLD_TITLE = u'3 坡地有限元模型建立与验证'
NEW_TITLE = u'3 斜入射成层坡地有限元模型与可信性验证'
NEXT_TITLE = u'4 坡地地震动放大效应参数化分析'


def file_sha256(path):
    """计算文件 SHA-256，确认源文件未被修改。"""
    h = hashlib.sha256()
    with open(path, 'rb') as fh:
        while True:
            chunk = fh.read(1024 * 1024)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest()


def element_text(element):
    """读取段落 XML 中的完整文本。"""
    return u''.join((node.text or u'') for node in element.iter(qn('w:t'))).strip()


def locate_chapter(body):
    """定位第三章标题和第四章标题对应的 body 子元素索引。"""
    children = list(body)
    start = end = None
    for idx, child in enumerate(children):
        text = element_text(child)
        if text == OLD_TITLE or text == NEW_TITLE:
            start = idx
        elif text == NEXT_TITLE and start is not None:
            end = idx
            break
    if start is None or end is None or end <= start:
        raise RuntimeError(u'无法定位第三章或第四章边界')
    return start, end


def replace_paragraph_text(paragraph_element, text):
    """保留段落属性，仅替换标题文本。"""
    for child in list(paragraph_element):
        if child.tag != qn('w:pPr'):
            paragraph_element.remove(child)
    run = OxmlElement('w:r')
    text_node = OxmlElement('w:t')
    text_node.text = text
    run.append(text_node)
    paragraph_element.append(run)


def make_paragraph(text, style_id, placeholder=False):
    """创建带指定 Word 样式的段落 XML。"""
    paragraph = OxmlElement('w:p')
    properties = OxmlElement('w:pPr')
    style = OxmlElement('w:pStyle')
    style.set(qn('w:val'), style_id)
    properties.append(style)
    paragraph.append(properties)

    run = OxmlElement('w:r')
    if placeholder:
        run_properties = OxmlElement('w:rPr')
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '808080')
        run_properties.append(color)
        run.append(run_properties)
    text_node = OxmlElement('w:t')
    text_node.set(qn('xml:space'), 'preserve')
    text_node.text = text
    run.append(text_node)
    paragraph.append(run)
    return paragraph


def archive_chapter():
    """在保留模板样式的副本中只留下重构前第三章。"""
    shutil.copy2(SOURCE, ARCHIVE_CH3)
    document = Document(ARCHIVE_CH3)
    body = document._element.body
    start, end = locate_chapter(body)
    children = list(body)
    keep = set(children[start:end])
    for child in children:
        if child.tag == qn('w:sectPr'):
            continue
        if child not in keep:
            body.remove(child)
    document.save(ARCHIVE_CH3)


def rebuild_working_copy():
    """生成新的第三章骨架，同时保留其他章节和全套模板格式。"""
    shutil.copy2(SOURCE, WORKING)
    document = Document(WORKING)
    body = document._element.body
    start, end = locate_chapter(body)
    children = list(body)
    start_element = children[start]
    end_element = children[end]

    replace_paragraph_text(start_element, NEW_TITLE)
    for child in children[start + 1:end]:
        body.remove(child)

    styles = document.styles
    heading2 = 'Heading2'  # 模板的二级标题样式 ID
    heading3 = 'Heading3'  # 模板的三级标题样式 ID
    body_text = 'BodyText'  # 模板的中文正文样式 ID

    blocks = [
        ('h2', u'3.1 研究目标、验证问题与适用边界'),
        ('p', u'【重构占位】P0/F0 完成后写入本章研究问题、证据等级和二维线性适用边界。'),
        ('h2', u'3.2 控制方程、基本假定与模型范围'),
        ('p', u'【重构占位】F0 与 V1 通过后写入控制方程、平面应变、材料本构、入射类型和结论边界。'),
        ('h2', u'3.3 有限元模型与数据协议'),
        ('h3', u'3.3.1 无量纲几何与成层介质'),
        ('h3', u'3.3.2 单元、网格、分析步和输出'),
        ('h3', u'3.3.3 斜入射等效输入与人工边界'),
        ('h3', u'3.3.4 阻尼、观测坐标与响应指标'),
        ('p', u'【重构占位】V1 通过后依次写入模型实现、参数表、模型示意和数据协议。'),
        ('h2', u'3.4 可信性框架、工况选择与评价准则'),
        ('h3', u'3.4.1 实现验证、数值解验证与适用性的区分'),
        ('h3', u'3.4.2 代表性和控制性工况选择'),
        ('h3', u'3.4.3 误差指标、事前门槛与误差预算'),
        ('p', u'【重构占位】V2 启动前写入已冻结的工况矩阵和验收判据。'),
        ('h2', u'3.5 端到端实现验证'),
        ('h3', u'3.5.1 几何、材料与观测生成检查'),
        ('h3', u'3.5.2 均质半空间解析解对比'),
        ('h3', u'3.5.3 平坦成层场地独立一维对比'),
        ('h3', u'3.5.4 方向镜像、同材退化与边界反射'),
        ('p', u'【重构占位】仅在 V1—V4 全部对应门槛通过后写入结果、误差图和结论。'),
        ('h2', u'3.6 数值解收敛与误差控制'),
        ('h3', u'3.6.1 控制工况与最短波长'),
        ('h3', u'3.6.2 网格、时间步和单元类型'),
        ('h3', u'3.6.3 计算域、观测窗和边界影响'),
        ('h3', u'3.6.4 阻尼、能量与最终误差预算'),
        ('p', u'【重构占位】V5 通过后写入四档收敛、GCI/误差预算及最终生产参数。'),
        ('h2', u'3.7 独立二维交叉验证与公开基准'),
        ('h3', u'3.7.1 独立二维方法或跨求解器对比'),
        ('h3', u'3.7.2 Shen 2024 均质坡地代表工况'),
        ('h3', u'3.7.3 Shen 2025 成层坡地代表工况'),
        ('h3', u'3.7.4 统一误差指标与外部一致性结论'),
        ('p', u'【重构占位】V6—V7 通过后写入量化对比；禁止只用云图相似或定性趋势。'),
        ('h2', u'3.8 可复现性、适用范围与局限'),
        ('h3', u'3.8.1 配置、数据、脚本哈希与回归集'),
        ('h3', u'3.8.2 线性幅值缩放与非线性决策边界'),
        ('h3', u'3.8.3 二维、频带、材料与工程外推限制'),
        ('p', u'【重构占位】V8 决策完成后写入线性适用域；涉及真实强震时必须补充 EQL/非线性敏感性。'),
        ('h2', u'3.9 本章小结'),
        ('p', u'【重构占位】V9 全回归通过后，只总结已经获得证据支持的平台能力。')
    ]

    for kind, text in blocks:
        style_id = heading2 if kind == 'h2' else heading3 if kind == 'h3' else body_text
        end_element.addprevious(make_paragraph(text, style_id, placeholder=(kind == 'p')))

    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(u'第3章，坡地有限元模型建立与验证。'):
            paragraph.text = (
                u'第3章，斜入射成层坡地有限元模型与可信性验证。区分实现验证、数值解验证、'
                u'独立二维基准和适用性边界，通过均质半空间解析对比、平坦成层独立一维对比、'
                u'多控制工况收敛、人工边界反射、跨方法与公开文献基准建立可复现证据链。'
            )
            paragraph.style = next(style for style in styles if style.style_id == body_text)
            break

    document.save(WORKING)


def validate_outputs(source_hash):
    """检查归档与工作稿的章节边界和旧结论清理状态。"""
    if file_sha256(SOURCE) != source_hash:
        raise RuntimeError(u'源修订稿哈希发生变化')

    archived = Document(ARCHIVE_CH3)
    archived_text = u'\n'.join(p.text for p in archived.paragraphs)
    if OLD_TITLE not in archived_text or NEXT_TITLE in archived_text:
        raise RuntimeError(u'第三章独立归档的章节边界不正确')

    working = Document(WORKING)
    working_text = u'\n'.join(p.text for p in working.paragraphs)
    for required in (NEW_TITLE, u'3.5.2 均质半空间解析解对比', u'3.9 本章小结', NEXT_TITLE):
        if required not in working_text:
            raise RuntimeError(u'工作稿缺少章节骨架: %s' % required)

    start = working_text.index(NEW_TITLE)
    end = working_text.index(NEXT_TITLE, start)
    chapter_text = working_text[start:end]
    forbidden = (u'522组', u'复现得到的522组', u'8 m网格已经达到', u'既有算例中该偏差均在2%以内')
    for phrase in forbidden:
        if phrase in chapter_text:
            raise RuntimeError(u'工作稿第三章仍含旧结论: %s' % phrase)

    print(u'SOURCE_SHA256=%s' % source_hash)
    print(u'ARCHIVE_CH3=%s' % ARCHIVE_CH3)
    print(u'ARCHIVE_CH3_SHA256=%s' % file_sha256(ARCHIVE_CH3))
    print(u'WORKING=%s' % WORKING)
    print(u'WORKING_SHA256=%s' % file_sha256(WORKING))
    print(u'WORKING_PARAGRAPHS=%d TABLES=%d' % (len(working.paragraphs), len(working.tables)))


def main():
    """执行归档、重构和结构校验。"""
    if not os.path.isfile(SOURCE):
        raise RuntimeError(u'找不到源修订稿: %s' % SOURCE)
    archive_dir = os.path.dirname(ARCHIVE_CH3)
    if not os.path.isdir(archive_dir):
        os.makedirs(archive_dir)
    source_hash = file_sha256(SOURCE)
    archive_chapter()
    rebuild_working_copy()
    validate_outputs(source_hash)


if __name__ == '__main__':
    main()
