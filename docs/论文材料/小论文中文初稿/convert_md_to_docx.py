# -*- coding: utf-8 -*-
"""
小论文中文初稿转换为投稿模板 Word 文档脚本
按照《投稿模板（中文稿件）.docx》规范生成《小论文中文初稿.docx》
"""

import os
import re
import sys
import copy
import zipfile
import shutil
import tempfile
import subprocess
from PIL import Image

import docx
from docx.shared import Pt, Inches, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn


def set_run_font(run, font_name_ascii='Times New Roman', font_name_eastasia='宋体', size_pt=10.5, bold=False, italic=False, superscript=False):
    """设置 Run 对象的西文和中文字体属性"""
    run.font.name = font_name_ascii
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if superscript:
        run.font.superscript = True
    
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name_ascii)
    rFonts.set(qn('w:hAnsi'), font_name_ascii)
    rFonts.set(qn('w:eastAsia'), font_name_eastasia)
    rFonts.set(qn('w:cs'), font_name_ascii)
    rFonts.set(qn('w:hint'), 'eastAsia')


def set_three_line_table(table, col_widths=None):
    """为表格设置学术期刊标准三线表边框与内边距"""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    
    new_borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="auto"/>\n'
        f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="auto"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:insideH w:val="none"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(new_borders)
    
    # 表头下细线
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is not None:
                tcPr.remove(tcBorders)
            new_tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>\n'
                f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="auto"/>\n'
                f'</w:tcBorders>'
            )
            tcPr.append(new_tcBorders)
            
    # 设置单元格宽度
    if col_widths and len(col_widths) == len(table.columns):
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = w


def batch_compile_math(latex_list):
    """
    通过 Pandoc 批量预编译所有 LaTeX 数学公式为 OMML 元素
    """
    unique_math = list(dict.fromkeys(latex_list))
    if not unique_math:
        return {}
    
    with tempfile.TemporaryDirectory() as tmpdir:
        md_file = os.path.join(tmpdir, "batch_math.md")
        docx_file = os.path.join(tmpdir, "batch_math.docx")
        
        with open(md_file, "w", encoding="utf-8") as f:
            for idx, m in enumerate(unique_math):
                clean_m = re.sub(r'\\tag\{.*?\}', '', m).strip()
                f.write(f"ITEM_{idx:04d}: ${clean_m}$\n\n")
        
        cmd = ["pandoc", md_file, "-o", docx_file]
        subprocess.run(cmd, capture_output=True, text=True)
        
        cache = {}
        if os.path.exists(docx_file):
            doc = docx.Document(docx_file)
            for p in doc.paragraphs:
                text = p.text
                m_match = re.search(r'ITEM_(\d+):', text)
                if m_match:
                    idx = int(m_match.group(1))
                    if idx < len(unique_math):
                        omaths = p._p.xpath('.//m:oMath')
                        if omaths:
                            cache[unique_math[idx]] = omaths[0]
        return cache


def batch_compile_display_math(equations):
    """
    批量编译独立显示公式为 OMML
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        md_file = os.path.join(tmpdir, "disp_math.md")
        docx_file = os.path.join(tmpdir, "disp_math.docx")
        
        with open(md_file, "w", encoding="utf-8") as f:
            for idx, (eq_tex, tag) in enumerate(equations):
                clean_tex = re.sub(r'\\tag\{.*?\}', '', eq_tex).strip()
                f.write(f"$${clean_tex}$$\n\n")
                
        cmd = ["pandoc", md_file, "-o", docx_file]
        subprocess.run(cmd, capture_output=True, text=True)
        
        omath_list = []
        if os.path.exists(docx_file):
            doc = docx.Document(docx_file)
            for p in doc.paragraphs:
                omaths = p._p.xpath('.//m:oMath')
                if omaths:
                    omath_list.append(omaths[0])
        return omath_list


def add_formatted_text(paragraph, text, math_cache, default_font_eastasia='宋体', default_font_ascii='Times New Roman', size_pt=10.5, default_bold=False, default_italic=False):
    """
    解析段落内的行内数学公式、上标、加粗等富文本格式并写入段落
    """
    pattern = r'(\$.*?\$|<sup>.*?</sup>|\*\*.*?\*\*)'
    tokens = re.split(pattern, text)
    
    for token in tokens:
        if not token:
            continue
        if token.startswith('$') and token.endswith('$') and len(token) >= 2:
            latex = token[1:-1]
            if latex in math_cache:
                omath_copy = copy.deepcopy(math_cache[latex])
                paragraph._p.append(omath_copy)
            else:
                r = paragraph.add_run(latex)
                set_run_font(r, font_name_ascii=default_font_ascii, font_name_eastasia=default_font_eastasia, size_pt=size_pt, bold=default_bold, italic=True)
        elif token.startswith('<sup>') and token.endswith('</sup>'):
            sub_text = token[5:-6]
            r = paragraph.add_run(sub_text)
            set_run_font(r, font_name_ascii=default_font_ascii, font_name_eastasia=default_font_eastasia, size_pt=size_pt, bold=default_bold, italic=default_italic, superscript=True)
        elif token.startswith('**') and token.endswith('**'):
            sub_text = token[2:-2]
            r = paragraph.add_run(sub_text)
            set_run_font(r, font_name_ascii=default_font_ascii, font_name_eastasia=default_font_eastasia, size_pt=size_pt, bold=True, italic=default_italic)
        else:
            r = paragraph.add_run(token)
            set_run_font(r, font_name_ascii=default_font_ascii, font_name_eastasia=default_font_eastasia, size_pt=size_pt, bold=default_bold, italic=default_italic)


def is_table_caption_line(line):
    s = line.strip()
    return bool(re.match(r'^表\s*\d+[\s\u3000\t]+', s) or re.match(r'^(?:Table|TABLE)\s*\d+[\s\u3000\t\.:]', s))


def is_figure_caption_line(line):
    s = line.strip()
    return bool(re.match(r'^图\s*\d+[\s\u3000\t]+', s) or re.match(r'^(?:Fig|Figure|FIG|FIGURE)\.?\s*\d+[\s\u3000\t\.:]', s))


def parse_markdown_blocks(md_content):
    """
    解析 Markdown 初稿为结构化数据块
    """
    lines = md_content.split('\n')
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        
        if line.startswith('# '):
            blocks.append(('TITLE_CN', line[2:].strip()))
            i += 1
        elif line.startswith('## '):
            blocks.append(('H1', line[3:].strip()))
            i += 1
        elif line.startswith('### '):
            blocks.append(('H2', line[4:].strip()))
            i += 1
        elif line.startswith('#### '):
            blocks.append(('H3', line[5:].strip()))
            i += 1
        elif line.startswith('$$'):
            math_lines = [line[2:].strip()]
            i += 1
            while i < len(lines) and not lines[i].strip().endswith('$$'):
                math_lines.append(lines[i].strip())
                i += 1
            if i < len(lines):
                end_line = lines[i].strip()
                if end_line != '$$':
                    math_lines.append(end_line[:-2].strip())
                i += 1
            blocks.append(('MATH_DISPLAY', "\n".join(math_lines)))
        elif line.startswith('!['):
            img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            img_path = img_match.group(2) if img_match else line
            i += 1
            captions = []
            while i < len(lines):
                s_line = lines[i].strip()
                if not s_line:
                    if i + 1 < len(lines) and is_figure_caption_line(lines[i+1]):
                        i += 1
                        continue
                    else:
                        break
                if s_line.startswith(('##', '###', '####', '![', '|', '$$')):
                    break
                if is_figure_caption_line(s_line) or (captions and not s_line.startswith(('表', 'Table'))):
                    captions.append(s_line)
                    i += 1
                else:
                    break
            blocks.append(('FIGURE', (img_path, captions)))
        elif is_table_caption_line(line) or line.startswith('|'):
            tbl_captions = []
            while i < len(lines) and not lines[i].strip().startswith('|'):
                s_line = lines[i].strip()
                if s_line:
                    if is_table_caption_line(s_line):
                        tbl_captions.append(s_line)
                    else:
                        break
                i += 1
            
            while i < len(lines) and not lines[i].strip().startswith('|') and not lines[i].strip():
                i += 1
                
            tbl_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl_rows.append(lines[i].strip())
                i += 1
                
            tbl_note = ""
            temp_i = i
            while temp_i < len(lines) and not lines[temp_i].strip():
                temp_i += 1
            if temp_i < len(lines) and lines[temp_i].strip().startswith('注：'):
                i = temp_i
                while i < len(lines) and lines[i].strip().startswith('注：'):
                    tbl_note += lines[i].strip()
                    i += 1
                    
            if tbl_rows:
                blocks.append(('TABLE', (tbl_captions, tbl_rows, tbl_note)))
            else:
                for cap in tbl_captions:
                    blocks.append(('PARA', cap))
        else:
            blocks.append(('PARA', line))
            i += 1
            
    return blocks


def update_docx_footnotes(docx_path):
    """更新 Word 文件中第一页脚注的内容为当前论文信息"""
    footnote_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:footnote w:type="separator" w:id="2">
    <w:p><w:r><w:separator/></w:r></w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="3">
    <w:p><w:r><w:continuationSeparator/></w:r></w:p>
  </w:footnote>
  <w:footnote w:id="0">
    <w:p>
      <w:pPr>
        <w:pStyle w:val="6"/>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体" w:cs="Times New Roman"/>
          <w:sz w:val="18"/>
        </w:rPr>
      </w:pPr>
      <w:r>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体" w:cs="Times New Roman"/>
          <w:sz w:val="18"/>
        </w:rPr>
        <w:t>收稿日期：2026-08-24</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:pPr>
        <w:pStyle w:val="6"/>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体" w:cs="Times New Roman"/>
          <w:sz w:val="18"/>
        </w:rPr>
      </w:pPr>
      <w:r>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体" w:cs="Times New Roman"/>
          <w:sz w:val="18"/>
        </w:rPr>
        <w:t>基金项目：国家自然科学基金(********)；重庆市自然科学基金(********)</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:pPr>
        <w:pStyle w:val="6"/>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体" w:cs="Times New Roman"/>
          <w:sz w:val="18"/>
        </w:rPr>
      </w:pPr>
      <w:r>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体" w:cs="Times New Roman"/>
          <w:sz w:val="18"/>
        </w:rPr>
        <w:t>作者简介：【作者一】（19XX- ），男，博士，副教授，主要从事边坡动力响应与数值模拟研究，E-mail：***@*******</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:pPr>
        <w:pStyle w:val="6"/>
        <w:ind w:firstLine="900" w:firstLineChars="500"/>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体" w:cs="Times New Roman"/>
          <w:sz w:val="18"/>
        </w:rPr>
      </w:pPr>
      <w:r>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体" w:cs="Times New Roman"/>
          <w:sz w:val="18"/>
        </w:rPr>
        <w:t>【作者二】（通信作者），男，教授，博士生导师，E-mail：***@*******。</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:pPr>
        <w:pStyle w:val="6"/>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体" w:cs="Times New Roman"/>
          <w:sz w:val="18"/>
        </w:rPr>
      </w:pPr>
      <w:r>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体" w:cs="Times New Roman"/>
          <w:sz w:val="18"/>
        </w:rPr>
        <w:t>Received: 2026-08-24</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:pPr>
        <w:pStyle w:val="6"/>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体" w:cs="Times New Roman"/>
          <w:sz w:val="18"/>
        </w:rPr>
      </w:pPr>
      <w:r>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体" w:cs="Times New Roman"/>
          <w:sz w:val="18"/>
        </w:rPr>
        <w:t>Foundation items: National Natural Science Foundation of China (No. ********); Natural Science Foundation of Chongqing (No. ********)</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:pPr>
        <w:pStyle w:val="6"/>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体" w:cs="Times New Roman"/>
          <w:sz w:val="18"/>
        </w:rPr>
      </w:pPr>
      <w:r>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体" w:cs="Times New Roman"/>
          <w:sz w:val="18"/>
        </w:rPr>
        <w:t>Author brief: 【AUTHOR 1】 (19XX- ), PhD, associate professor, main research interest: slope dynamic response and numerical simulation, E-mail: ***@*******</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:pPr>
        <w:pStyle w:val="6"/>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体" w:cs="Times New Roman"/>
          <w:sz w:val="18"/>
        </w:rPr>
      </w:pPr>
      <w:r>
        <w:rPr>
          <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体" w:cs="Times New Roman"/>
          <w:sz w:val="18"/>
        </w:rPr>
        <w:t>           【AUTHOR 2】 (corresponding author), professor, doctorial supervisor, ***@*******.</w:t>
      </w:r>
    </w:p>
  </w:footnote>
</w:footnotes>"""
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_docx = os.path.join(tmpdir, "temp.docx")
        with zipfile.ZipFile(docx_path, 'r') as zin:
            with zipfile.ZipFile(tmp_docx, 'w') as zout:
                for item in zin.infolist():
                    if item.filename == 'word/footnotes.xml':
                        zout.writestr(item, footnote_xml.encode('utf-8'))
                    else:
                        zout.writestr(item, zin.read(item.filename))
        shutil.copy2(tmp_docx, docx_path)


def build_word_document(md_path, template_path, output_path):
    """
    主构建函数：将 Markdown 初稿转换为规范 Word 稿件
    """
    base_dir = os.path.dirname(os.path.abspath(md_path))
    
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
        
    blocks = parse_markdown_blocks(md_content)
    
    # 过滤显示公式后再提取行内公式
    clean_for_inline = re.sub(r'\$\$.*?\$\$', '', md_content, flags=re.DOTALL)
    inline_math_list = re.findall(r'(?<!\$)\$(?!\$)(.*?)(?<!\$)\$(?!\$)', clean_for_inline)
    print(f"正在编译 {len(inline_math_list)} 个行内公式...")
    math_cache = batch_compile_math(inline_math_list)
    print(f"行内公式编译完成，缓存 {len(math_cache)} 个唯一公式。")
    
    # 提取显示公式
    display_equations = []
    for b in blocks:
        if b[0] == 'MATH_DISPLAY':
            eq_text = b[1]
            tag_match = re.search(r'\\tag\{(\d+)\}', eq_text)
            tag = f"({tag_match.group(1)})" if tag_match else ""
            display_equations.append((eq_text, tag))
            
    print(f"正在编译 {len(display_equations)} 个独立显示公式...")
    disp_omath_list = batch_compile_display_math(display_equations)
    print(f"显示公式编译完成，获得 {len(disp_omath_list)} 个 OMML 对象。")
    
    # 基于模板加载文档
    doc = docx.Document(template_path)
    
    # 清空模板中默认的 Body 段落与表格，但保留样式、节设置与页眉页脚
    body_elem = doc._body._element
    for child in list(body_elem):
        if child.tag.endswith(('p', 'tbl')):
            body_elem.remove(child)
            
    # 设置页眉文字
    header_text = "【作者一】，等：成层坡地地震动放大效应的幅值—相位联合表征、代理预测与真实地震动重构"
    for s in doc.sections:
        for p in s.header.paragraphs:
            p.text = header_text
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in p.runs:
                set_run_font(r, font_name_ascii='Times New Roman', font_name_eastasia='宋体', size_pt=9.0)
                
    # 分离文头部分与正文部分
    front_blocks = []
    body_blocks = []
    is_body = False
    
    for b in blocks:
        if b[0] == 'H1' and ('1 ' in b[1] or '引' in b[1]):
            is_body = True
        if is_body:
            body_blocks.append(b)
        else:
            front_blocks.append(b)
            
    print(f"解析获得文头块数：{len(front_blocks)}，正文块数：{len(body_blocks)}")
    
    # ==========================
    # 1. 构建第 1 节：单栏文头
    # ==========================
    # 中文标题
    title_cn = "成层坡地地震动放大效应的幅值—相位联合表征、代理预测与真实地震动重构"
    p_tcn = doc.add_paragraph()
    p_tcn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tcn.paragraph_format.space_before = Pt(23.4)
    p_tcn.paragraph_format.space_after = Pt(0)
    p_tcn.paragraph_format.left_indent = Pt(21)
    p_tcn.paragraph_format.right_indent = Pt(21)
    r_tcn = p_tcn.add_run(title_cn)
    set_run_font(r_tcn, font_name_ascii='Times New Roman', font_name_eastasia='宋体', size_pt=16.0, bold=True)
    
    # 挂载首页脚注引用
    r_fn = p_tcn.add_run()
    fn_ref = parse_xml(f'<w:footnoteReference {nsdecls("w")} w:id="0"/>')
    r_fn._r.append(fn_ref)
    
    # 中文作者
    p_acn = doc.add_paragraph()
    p_acn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_acn.paragraph_format.space_before = Pt(7.8)
    p_acn.paragraph_format.space_after = Pt(0)
    p_acn.paragraph_format.left_indent = Pt(21)
    p_acn.paragraph_format.right_indent = Pt(21)
    add_formatted_text(p_acn, "【作者一】<sup>1</sup>，【作者二】<sup>1</sup>，【作者三】<sup>1</sup>", math_cache, default_font_eastasia='仿宋', default_font_ascii='Times New Roman', size_pt=12.0)
    
    # 中文单位
    p_afcn = doc.add_paragraph()
    p_afcn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_afcn.paragraph_format.space_before = Pt(0)
    p_afcn.paragraph_format.space_after = Pt(7.8)
    p_afcn.paragraph_format.left_indent = Pt(21)
    p_afcn.paragraph_format.right_indent = Pt(21)
    add_formatted_text(p_afcn, "（1. 【单位全称，省 市 邮编】）", math_cache, default_font_eastasia='宋体', default_font_ascii='Times New Roman', size_pt=10.5)
    
    # 中文摘要
    abstract_cn = ""
    keywords_cn = ""
    for b in front_blocks:
        if b[0] == 'PARA' and ('摘 要：' in b[1] or '摘要：' in b[1]):
            abstract_cn = b[1]
        elif b[0] == 'PARA' and ('关 键 词：' in b[1] or '关键词：' in b[1]):
            keywords_cn = b[1]
            
    if abstract_cn:
        p_abscn = doc.add_paragraph()
        p_abscn.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_abscn.paragraph_format.left_indent = Pt(21)
        p_abscn.paragraph_format.right_indent = Pt(21)
        if '：' in abstract_cn:
            prefix, content = abstract_cn.split('：', 1)
            prefix = "摘  要："
        else:
            prefix, content = "摘  要：", abstract_cn
        r_lbl = p_abscn.add_run(prefix)
        set_run_font(r_lbl, font_name_ascii='Times New Roman', font_name_eastasia='黑体', size_pt=10.5, bold=False)
        add_formatted_text(p_abscn, content, math_cache, default_font_eastasia='楷体', default_font_ascii='Times New Roman', size_pt=10.5)
        
    # 中文关键词
    if keywords_cn:
        p_kwcn = doc.add_paragraph()
        p_kwcn.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_kwcn.paragraph_format.left_indent = Pt(21)
        p_kwcn.paragraph_format.right_indent = Pt(21)
        if '：' in keywords_cn:
            prefix, content = keywords_cn.split('：', 1)
            prefix = "关键词："
        else:
            prefix, content = "关键词：", keywords_cn
        r_lbl = p_kwcn.add_run(prefix)
        set_run_font(r_lbl, font_name_ascii='Times New Roman', font_name_eastasia='黑体', size_pt=10.5, bold=False)
        add_formatted_text(p_kwcn, content, math_cache, default_font_eastasia='楷体', default_font_ascii='Times New Roman', size_pt=10.5)
        
    # 中图分类号与文献标志码
    p_clc = doc.add_paragraph()
    p_clc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_clc.paragraph_format.left_indent = Pt(21)
    p_clc.paragraph_format.right_indent = Pt(21)
    r_clc1 = p_clc.add_run("中图分类号：")
    set_run_font(r_clc1, font_name_ascii='Times New Roman', font_name_eastasia='黑体', size_pt=10.5)
    r_clc2 = p_clc.add_run("TU43                ")
    set_run_font(r_clc2, font_name_ascii='Times New Roman', font_name_eastasia='宋体', size_pt=10.5)
    r_clc3 = p_clc.add_run("文献标志码：")
    set_run_font(r_clc3, font_name_ascii='Times New Roman', font_name_eastasia='黑体', size_pt=10.5)
    r_clc4 = p_clc.add_run("A")
    set_run_font(r_clc4, font_name_ascii='Times New Roman', font_name_eastasia='Times New Roman', size_pt=10.5)
    
    # 英文题目
    title_en = "Amplitude–phase joint characterization, surrogate prediction and ground-motion reconstruction of seismic amplification effects of layered slopes"
    p_ten = doc.add_paragraph()
    p_ten.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ten.paragraph_format.space_before = Pt(15.6)
    p_ten.paragraph_format.space_after = Pt(0)
    r_ten = p_ten.add_run(title_en)
    set_run_font(r_ten, font_name_ascii='Times New Roman', font_name_eastasia='Times New Roman', size_pt=14.0, bold=True)
    
    # 英文作者
    p_aen = doc.add_paragraph()
    p_aen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_aen.paragraph_format.space_before = Pt(7.8)
    p_aen.paragraph_format.space_after = Pt(0)
    add_formatted_text(p_aen, "【AUTHOR 1】<sup>1</sup>, 【AUTHOR 2】<sup>1</sup>, 【AUTHOR 3】<sup>1</sup>", math_cache, default_font_eastasia='Times New Roman', default_font_ascii='Times New Roman', size_pt=12.0)
    
    # 英文单位
    p_afen = doc.add_paragraph()
    p_afen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_afen.paragraph_format.space_before = Pt(0)
    p_afen.paragraph_format.space_after = Pt(7.8)
    add_formatted_text(p_afen, "(1. 【Affiliation, City, Postcode, China】)", math_cache, default_font_eastasia='Times New Roman', default_font_ascii='Times New Roman', size_pt=10.5)
    
    # 英文摘要与关键词
    abstract_en = ""
    keywords_en = ""
    for b in front_blocks:
        if b[0] == 'PARA' and ('Abstract:' in b[1] or 'Abstract：' in b[1]):
            abstract_en = b[1]
        elif b[0] == 'PARA' and ('Key words:' in b[1] or 'Keywords:' in b[1]):
            keywords_en = b[1]
            
    if abstract_en:
        p_absen = doc.add_paragraph()
        p_absen.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if 'Abstract:' in abstract_en:
            prefix, content = abstract_en.split('Abstract:', 1)
            prefix = "Abstract: "
        elif 'Abstract：' in abstract_en:
            prefix, content = abstract_en.split('Abstract：', 1)
            prefix = "Abstract: "
        else:
            prefix, content = "Abstract: ", abstract_en
        r_lbl = p_absen.add_run(prefix)
        set_run_font(r_lbl, font_name_ascii='Times New Roman', font_name_eastasia='Times New Roman', size_pt=10.5, bold=True)
        add_formatted_text(p_absen, content.strip(), math_cache, default_font_eastasia='Times New Roman', default_font_ascii='Times New Roman', size_pt=10.5)
        
    if keywords_en:
        p_kwen = doc.add_paragraph()
        p_kwen.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_kwen.paragraph_format.space_after = Pt(6.0)
        if 'Key words:' in keywords_en:
            prefix, content = keywords_en.split('Key words:', 1)
            prefix = "Keywords: "
        elif 'Keywords:' in keywords_en:
            prefix, content = keywords_en.split('Keywords:', 1)
            prefix = "Keywords: "
        else:
            prefix, content = "Keywords: ", keywords_en
        r_lbl = p_kwen.add_run(prefix)
        set_run_font(r_lbl, font_name_ascii='Times New Roman', font_name_eastasia='Times New Roman', size_pt=10.5, bold=True)
        add_formatted_text(p_kwen, content.strip(), math_cache, default_font_eastasia='Times New Roman', default_font_ascii='Times New Roman', size_pt=10.5)
        
    # 文头结束：添加连续分节符，进入双栏正文
    p_last_front = doc.paragraphs[-1]
    pPr_front = p_last_front._p.get_or_add_pPr()
    sectPr_front = parse_xml(
        f'<w:sectPr {nsdecls("w")}>\n'
        f'  <w:type w:val="continuous"/>\n'
        f'  <w:pgSz w:w="11906" w:h="16838"/>\n'
        f'  <w:pgMar w:top="1418" w:right="1134" w:bottom="1418" w:left="1134" w:header="851" w:footer="992" w:gutter="0"/>\n'
        f'  <w:cols w:num="1" w:space="425"/>\n'
        f'  <w:docGrid w:type="lines" w:linePitch="312"/>\n'
        f'</w:sectPr>'
    )
    pPr_front.append(sectPr_front)
    
    # ==========================
    # 2. 构建正文部分（双栏排版）
    # ==========================
    disp_math_idx = 0
    in_references = False
    
    for b_idx, b in enumerate(body_blocks):
        b_type = b[0]
        
        if b_type == 'H1':
            h1_text = b[1]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(15.6)
            p.paragraph_format.space_after = Pt(3.0)
            p.paragraph_format.first_line_indent = Pt(0)
            r = p.add_run(h1_text)
            set_run_font(r, font_name_ascii='Times New Roman', font_name_eastasia='宋体', size_pt=14.0, bold=True)
            if '参考' in h1_text:
                in_references = True
                
        elif b_type == 'H2':
            h2_text = b[1]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(7.8)
            p.paragraph_format.space_after = Pt(2.0)
            p.paragraph_format.first_line_indent = Pt(0)
            r = p.add_run(h2_text)
            set_run_font(r, font_name_ascii='Times New Roman', font_name_eastasia='宋体', size_pt=12.0, bold=True)
            
        elif b_type == 'H3':
            h3_text = b[1]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(4.0)
            p.paragraph_format.space_after = Pt(1.0)
            p.paragraph_format.first_line_indent = Pt(0)
            r = p.add_run(h3_text)
            set_run_font(r, font_name_ascii='Times New Roman', font_name_eastasia='宋体', size_pt=10.5, bold=True)
            
        elif b_type == 'MATH_DISPLAY':
            if disp_math_idx < len(disp_omath_list):
                omath_elem = disp_omath_list[disp_math_idx]
                tag_str = display_equations[disp_math_idx][1]
                disp_math_idx += 1
                
                eq_tbl = doc.add_table(rows=1, cols=2)
                eq_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                eq_tbl.autofit = False
                
                # 单栏版心下公式表总宽约6.6in，公式列居中、编号列贴右侧
                eq_tbl.columns[0].width = Inches(5.9)
                eq_tbl.columns[1].width = Inches(0.7)
                
                tblPr = eq_tbl._tbl.tblPr
                tblBorders = parse_xml(
                    f'<w:tblBorders {nsdecls("w")}>\n'
                    f'  <w:top w:val="none"/>\n'
                    f'  <w:bottom w:val="none"/>\n'
                    f'  <w:left w:val="none"/>\n'
                    f'  <w:right w:val="none"/>\n'
                    f'  <w:insideH w:val="none"/>\n'
                    f'  <w:insideV w:val="none"/>\n'
                    f'</w:tblBorders>'
                )
                tblPr.append(tblBorders)
                
                c0 = eq_tbl.cell(0, 0)
                p0 = c0.paragraphs[0]
                p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p0.paragraph_format.space_before = Pt(3.0)
                p0.paragraph_format.space_after = Pt(3.0)
                p0._p.append(copy.deepcopy(omath_elem))
                
                c1 = eq_tbl.cell(0, 1)
                p1 = c1.paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p1.paragraph_format.space_before = Pt(3.0)
                p1.paragraph_format.space_after = Pt(3.0)
                r_tag = p1.add_run(tag_str)
                set_run_font(r_tag, font_name_ascii='Times New Roman', font_name_eastasia='Times New Roman', size_pt=10.5)
                
        elif b_type == 'FIGURE':
            img_rel_path, captions = b[1]
            img_full_path = os.path.join(base_dir, img_rel_path)
            
            if os.path.exists(img_full_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(4.0)
                p_img.paragraph_format.space_after = Pt(2.0)
                p_img.paragraph_format.first_line_indent = Pt(0)
                
                im = Image.open(img_full_path)
                w, h = im.size
                # 单栏版心约6.69in，图宽取6.2in，同时限制图高不超过6.5in
                target_w = min(Inches(6.2), Inches(6.5) * w / h)
                p_img.add_run().add_picture(img_full_path, width=target_w)
                
                for c_idx, cap in enumerate(captions):
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.first_line_indent = Pt(0)
                    p_cap.paragraph_format.space_before = Pt(0)
                    p_cap.paragraph_format.space_after = Pt(3.0 if c_idx == len(captions)-1 else 0)
                    
                    is_en = cap.startswith(('Fig', 'fig'))
                    eastasia_font = 'Times New Roman' if is_en else '宋体'
                    add_formatted_text(p_cap, cap, math_cache, default_font_eastasia=eastasia_font, default_font_ascii='Times New Roman', size_pt=9.0)
            else:
                print(f"警告：找不到图片文件 {img_full_path}")
                
        elif b_type == 'TABLE':
            tbl_captions, tbl_rows, tbl_note = b[1]
            
            for c_idx, cap in enumerate(tbl_captions):
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.first_line_indent = Pt(0)
                p_cap.paragraph_format.space_before = Pt(7.8 if c_idx == 0 else 0)
                p_cap.paragraph_format.space_after = Pt(2.0 if c_idx == len(tbl_captions)-1 else 0)
                
                is_en = cap.startswith(('Table', 'table'))
                eastasia_font = 'Times New Roman' if is_en else '宋体'
                add_formatted_text(p_cap, cap, math_cache, default_font_eastasia=eastasia_font, default_font_ascii='Times New Roman', size_pt=9.0)
                
            parsed_rows = []
            for r_line in tbl_rows:
                if re.match(r'^\s*\|(?:\s*:?-+:?\s*\|)+\s*$', r_line):
                    continue
                cells = [c.strip() for c in r_line.strip('|').split('|')]
                parsed_rows.append(cells)
                
            if parsed_rows:
                num_rows = len(parsed_rows)
                num_cols = max(len(r) for r in parsed_rows)
                
                tbl = doc.add_table(rows=num_rows, cols=num_cols)
                set_three_line_table(tbl)
                
                for r_i, r_data in enumerate(parsed_rows):
                    is_header = (r_i == 0)
                    for c_i, cell_text in enumerate(r_data):
                        if c_i < num_cols:
                            cell = tbl.cell(r_i, c_i)
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.paragraph_format.space_before = Pt(1.5)
                            p.paragraph_format.space_after = Pt(1.5)
                            p.paragraph_format.first_line_indent = Pt(0)
                            
                            font_size = 8.0 if num_cols >= 7 else 9.0
                            add_formatted_text(p, cell_text, math_cache, default_font_eastasia='宋体', default_font_ascii='Times New Roman', size_pt=font_size, default_bold=is_header)
                            
            if tbl_note:
                p_note = doc.add_paragraph()
                p_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_note.paragraph_format.first_line_indent = Pt(18)
                p_note.paragraph_format.space_before = Pt(2.0)
                p_note.paragraph_format.space_after = Pt(4.0)
                add_formatted_text(p_note, tbl_note, math_cache, default_font_eastasia='宋体', default_font_ascii='Times New Roman', size_pt=8.5)
                
        elif b_type == 'PARA':
            para_text = b[1]
            p = doc.add_paragraph()
            
            if in_references:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2.0)
                
                if para_text.startswith('['):
                    p.paragraph_format.left_indent = Pt(13.5)
                    p.paragraph_format.first_line_indent = Pt(-13.5)
                    add_formatted_text(p, para_text, math_cache, default_font_eastasia='宋体', default_font_ascii='Times New Roman', size_pt=9.0)
                else:
                    p.paragraph_format.left_indent = Pt(15.7)
                    p.paragraph_format.first_line_indent = Pt(0)
                    add_formatted_text(p, para_text, math_cache, default_font_eastasia='Times New Roman', default_font_ascii='Times New Roman', size_pt=9.0)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Pt(21)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                add_formatted_text(p, para_text, math_cache, default_font_eastasia='宋体', default_font_ascii='Times New Roman', size_pt=10.5)
                
    # 正文采用单栏排版
    doc_sectPr = doc.sections[-1]._sectPr
    doc_cols = doc_sectPr.find(qn('w:cols'))
    if doc_cols is None:
        doc_cols = OxmlElement('w:cols')
        doc_sectPr.append(doc_cols)
    doc_cols.set(qn('w:num'), '1')
    
    # 保存生成的文档
    doc.save(output_path)
    
    # 更新首页脚注内容
    update_docx_footnotes(output_path)
    print(f"Word 文档已成功生成并保存至：{output_path}")


if __name__ == '__main__':
    # 基于脚本自身定位初稿目录，不依赖运行时工作目录
    default_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = sys.argv[1] if len(sys.argv) > 1 else os.path.join(default_dir, '小论文中文初稿.md')
    md_file = os.path.abspath(md_file)
    
    if len(sys.argv) > 2:
        out_file = os.path.abspath(sys.argv[2])
    else:
        out_file = os.path.splitext(md_file)[0] + '.docx'
        
    if len(sys.argv) > 3:
        tpl_file = os.path.abspath(sys.argv[3])
    else:
        tpl_file = os.path.join(default_dir, '投稿模板（中文稿件）.docx')
    
    print(f"正在从 Markdown 构建 Word 稿件...")
    print(f"输入文件: {md_file}")
    print(f"模板文件: {tpl_file}")
    print(f"输出文件: {out_file}")
    
    build_word_document(md_file, tpl_file, out_file)
