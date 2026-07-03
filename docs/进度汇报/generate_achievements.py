import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def set_run_font(run, size_pt=10.5, bold=False, italic=False, color_rgb=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color_rgb:
        run.font.color.rgb = color_rgb
    run.font.name = 'Calibri'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_heading_1(doc, text):
    p = doc.add_heading(level=1)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size_pt=15, bold=True, color_rgb=RGBColor(31, 78, 121))
    return p

def add_heading_2(doc, text):
    p = doc.add_heading(level=2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size_pt=12, bold=True, color_rgb=RGBColor(31, 78, 121))
    return p

def add_heading_3(doc, text):
    p = doc.add_heading(level=3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size_pt=11, bold=True, color_rgb=RGBColor(31, 78, 121))
    return p

def add_normal_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size_pt=10.5, bold=bold)
    return p

def add_bullet_paragraph(doc, prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    run_pref = p.add_run(prefix)
    set_run_font(run_pref, size_pt=10.5, bold=True)
    
    run_text = p.add_run(text)
    set_run_font(run_text, size_pt=10.5, bold=False)
    return p

def add_figure_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size_pt=9, bold=False, color_rgb=RGBColor(80, 80, 80))
    return p

def add_figure_placeholder(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shading_elm = parse_xml(r'<w:shd {} w:fill="FAFAFA"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"【图表占位框：{text}】")
    set_run_font(run, size_pt=9.5, bold=True, color_rgb=RGBColor(120, 120, 120))
    
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(2)
    p_space.paragraph_format.space_after = Pt(2)

def create_styled_table(doc, headers, data):
    table = doc.add_table(rows=1 + len(data), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表头
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        shading_elm = parse_xml(r'<w:shd {} w:fill="1F4E79"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
        
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.runs[0]
        set_run_font(run, size_pt=9.5, bold=True, color_rgb=RGBColor(255, 255, 255))
        
    # 数据行
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F9F9F9" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            if bg_color != "FFFFFF":
                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
                row_cells[c_idx]._tc.get_or_add_tcPr().append(shading_elm)
            
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if len(row_data) > 2 and c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if len(p.runs) > 0:
                run = p.runs[0]
                set_run_font(run, size_pt=9, bold=False)
                
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(2)
    p_space.paragraph_format.space_after = Pt(2)

def main():
    doc = Document()
    
    # A4 页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = docx.shared.Cm(2.54)
        section.bottom_margin = docx.shared.Cm(2.54)
        section.left_margin = docx.shared.Cm(2.54)
        section.right_margin = docx.shared.Cm(2.54)

    # 标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(14)
    title_p.paragraph_format.space_after = Pt(4)
    run = title_p.add_run("近期研究进展汇报")
    set_run_font(run, size_pt=18, bold=True, color_rgb=RGBColor(31, 78, 121))

    # 副标题
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_before = Pt(2)
    subtitle_p.paragraph_format.space_after = Pt(12)
    run_sub = subtitle_p.add_run("成层斜坡有限元数值动力模拟与机器学习代理模型研究\n2026年6月")
    set_run_font(run_sub, size_pt=11, bold=False, color_rgb=RGBColor(80, 80, 80))

    # ==================== PART I ====================
    add_heading_1(doc, "Part I · Abaqus 成层斜坡数值动力模拟进展")
    
    add_heading_2(doc, "1.1 研究背景与痛点")
    add_normal_paragraph(doc, "本研究针对地震波斜入射条件下多地层/成层坡地地形放大效应开展有限元数值模拟，采用 Abaqus 构建平面应变模型，设置粘弹性人工边界（VAB）并施加等效节点力。")
    add_normal_paragraph(doc, "在前期多地层数值模拟的尝试中，使用初始时域射线叠加法（v3-v5 脚本）计算得到的覆盖层地表地形放大系数（TAF）及峰值加速度（PGA）系统性偏小。通过深入剖析，发现了射线法在推广至多层场地时的物理局限，从而促成了波动输入算法由时域射线法向频域全局矩阵法的重大升级。")

    headers_summary = ["已跑工况数据 (E:\\Abaqus)", "无量纲参数维度", "v8 多层脚本规模", "V1-V4 对拍校验"]
    data_summary = [["522 组（均质单质坡）", "5 维 (η, i, θ, d/h, α)", "2479 行（含注释）", "100% 校验通过"]]
    create_styled_table(doc, headers_summary, data_summary)

    add_heading_2(doc, "1.2 波动输入算法升级：从射线法到矩阵法")
    
    add_bullet_paragraph(doc, "· 旧时域射线叠加法（v3-v5）的误差剖析：", "射线法将自由场近似为有限条折线波至的延迟叠加。在多层场地中，该方法引致了六大误差：① 界面 P-SV 转换波缺失；② 腔内多次波几何级数强行截断；③ 丢弃跨腔混响路径（腔间解耦）；④ 边界内外阻尼介质不一致；⑤ 也是最隐蔽的透反射系数应力与位移口径错配；⑥ 直达路径缺失单程透射放大因子。这些误差使得侧边界节点的自由场运动被系统性严重低估（幅值仅为精确解的几分之一），直接揭示了多层坡地动力计算结果偏小的输入侧成因。")
    
    add_bullet_paragraph(doc, "· 新频域全局矩阵法（v8 fd 引擎）的优势：", "基于 Thomson-Haskell 全局矩阵法，在频率域内对边界各节点所在的成层柱列出一维波动方程，批量一次性精确求解 4M+2 个未知波幅的线性方程组。该方法从物理上天然包揽了界面波形转换、各层任意多次往返混响、穿层耦合，并通过复模量和复密度引入瑞利阻尼，确保了自由场衰减与有限元域介质衰减的严格口径一致。")

    headers_compare = ["物理特性维度", "时域射线叠加（旧，v3-v5）", "频域全局矩阵（新，v8 fd引擎）"]
    data_compare = [
        ["均质半空间", "精确（解析走时延迟叠加）", "精确（逐项还原解析解对拍）"],
        ["界面透反射系数", "阻抗标量近似 + 应力/位移口径错配（低估幅值）", "Zoeppritz 级精确（隐含于方程矩阵）"],
        ["界面 SV-P 转换", "置零（丢弃转换 P 波）", "精确包含（求解多维极化向量）"],
        ["层内多次往返波", "几何级数截断（默认取 3 阶多次波）", "无穷阶混响全部包含，无截断误差"],
        ["跨层混响耦合", "忽略（各有限层视为独立腔体）", "精确包含（全局系统联立求解）"],
        ["边界介质阻尼衰减", "弹性自由场（与 FE 域瑞利阻尼失配）", "复模量复密度（与 FE 域阻尼严格同口径）"],
        ["离散化与参数敏感度", "敏感度高（结果随截断阶数漂移）", "无离散化参数，仅受 FFT 采样控制"],
        ["适用介质层数", "仅适用于 1-2 层，层数多时累积误差显著", "适用于任意层数场地，无层数限制"]
    ]
    create_styled_table(doc, headers_compare, data_compare)

    add_heading_2(doc, "1.3 多层斜坡有限元自洽建模实现")
    add_normal_paragraph(doc, "从均质单层升级为支持多层介质的 v8 版本，在 Abaqus 二维动力学有限元建模中，主要通过以下五项核心机制实现多层介质的精确几何剖分与物理自洽：")
    
    add_bullet_paragraph(doc, "1. 多层场地带（Band）构造与小面质心（Centroid）定位技术：", "脚本使用 _build_stratigraphy 函数将场地从下至上表示为虚拟材料带列表。在几何切分后，面临将多类材料赋予不同面（Face）的难题。本方案采用面质心定位算法（_assign_sections_by_band）：计算每个切分小面的几何质心坐标 (xc, yc)，调用 _band_bounds_at 自动换算该水平位置下各层界面的高程上下界，精确判断小面落入的材料带，从而为大量被切割的小面自动且精确地赋予对应的材料截面。")
    
    add_bullet_paragraph(doc, "2. 两类表层几何切分机制（Horizontal vs Terrain）：", "为满足不同边坡实际，脚本提供两类几何切分方式。① 水平分层模式（horizontal）：各有限层上下界为绝对水平线，在草图相应高程处画水平直线切割模型；② Follow 地形等厚铺设模式（terrain）：表层覆盖层沿地表起伏等厚铺设。通过计算由“上平台-坡面-下平台”三段折线构成的等深偏置线（Offset Line），利用 Abaqus 几何切割工具沿折线剖分模型，实现地层随地形起伏的等厚铺设。")
    
    add_bullet_paragraph(doc, "3. 粘弹性人工边界（VAB）弹簧-阻尼参数的逐层取材匹配：", "侧向人工边界截面上跨越了不同的材料层。为防止产生非物理的伪反射，方案引入边界节点局部取材技术（pick_material）。在边界弹簧/阻尼器生成阶段，遍历侧边界的所有节点，按节点坐标动态查询其落入的材料带，并提取该层介质的局部剪切模量 G、波速 cs/cp 和密度 ρ，逐节点单独计算刚度与阻力系数（kn, cn, kt, ct），实现吸收边界与局部介质阻抗的逐点精确匹配。")
    
    add_bullet_paragraph(doc, "4. 基于控制性“最软层”波速的 K-L 自适应网格控制：", "剪切波在软覆盖层中波长短、对网格分辨率要求高，而在坚硬基岩中波长长。脚本自动搜索各有限层中波速最低的“控制性最软层”剪切波速 cs_min，按 Kuhlemeyer-Lysmer 判据计算最大网格尺寸 Δl ≤ cs_min / (10 · f_max)（截止频率 f_max = 2.5 · fc），在全模型统一采用 CPE4R 二维减缩积分单元进行自适应网格划分，兼顾数值精度与计算规模。")
    
    add_bullet_paragraph(doc, "5. 频率相关瑞利阻尼的层间一致性标定：", "多层场地中各层介质的品质因子 Q 各不相同，阻尼比为 ξi = 1/(2Qi)。脚本采用双控锚定机制（anchor: 'dual'），拟合低频锚点 f1 = min(f1_factor · fc, f_site)（其中 f_site 为多层场地基频）与高频锚点 f2，独立计算各层的 Rayleigh 系数 (αi, βi)，在 Abaqus 材料中逐层赋予；同时在频域自由场（fd 引擎）中将阻尼以复模量方式精确引入，保证边界等效力和有限元内部介质的瑞利衰减严格一致。")

    add_heading_2(doc, "1.4 计算现状与下一步工作计划")
    add_bullet_paragraph(doc, "· 算法与脚本现状：", "多地层 v8 建模脚本及 fd 引擎已经过半空间退化、垂直入射 Haskell 递推、边界对拍 QA 等验证，确认支持多地层的数值计算无反射污染。")
    add_bullet_paragraph(doc, "· 工况计算现状：", "多层斜坡模型数据当前尚未批量运行。已经跑完的所有有限元工况数据全部保存在 E:\\Abaqus 目录下，均是使用**均匀单质地层初始脚本（v2 基线）**跑出的数据。")
    add_bullet_paragraph(doc, "· 下一步数值模拟计划：", "利用已调试通过的 v8 多地层脚本，启动 78 组多地层工况（含坡角、覆土层厚、介质阻抗比交叉组合）的有限元夜间批跑，提取其多层场地传递函数。")

    # ==================== PART II ====================
    add_heading_1(doc, "Part II · 机器学习代理模型研究进展")
    
    add_heading_2(doc, "2.1 数据来源澄清与数据集构成")
    add_normal_paragraph(doc, "本阶段机器学习旨在构建基于本征正交分解（POD）与高斯过程回归（GPR）的快速地表响应代理模型。")
    add_normal_paragraph(doc, "【核心说明】：本阶段机器学习所使用的数据来源全部是初始脚本跑的均质单质坡地有限元模型数据（即 E:\\Abaqus 目录下的 522 组样本）。多地层有限元数据尚在等待下一步计算，暂未纳入当前的机器学习模型。")

    headers_dataset = ["数据要素", "描述说明", "取值与物理量规模"]
    data_dataset = [
        ["总样本量", "均质工况几何 × 输入地震波", "174 × 3 = 522 组 TAF 响应曲线"],
        ["几何工况组合", "坡角 (i) × 坡高 (h) 交叉组合", "174 组均匀单质坡地几何"],
        ["输入地震波", "El Centro / Loma Prieta / Northridge", "3 条（PGA 调幅一致）"],
        ["测点长度", "沿坡面地表的空间取样点数", "161 个节点"],
        ["POD 降维阶数", "累计方差贡献率 ≥ 99.5%", "前 15 阶解耦 POD 系数"]
    ]
    create_styled_table(doc, headers_dataset, data_dataset)

    add_heading_2(doc, "2.2 方法框架（POD + GPR）")
    add_normal_paragraph(doc, "10 维输入特征：几何特征（坡角、坡高、坡脚距、坡顶宽）+ 地震波特征（PGA、卓越周期、反应谱 T=0.1/0.5/1.0s 的谱值 Sa）。")
    add_normal_paragraph(doc, "SVD 降维提取前 15 个正交基，将 161 维 TAF 曲线压缩至 15 维 POD 系数空间，对 15 个系数分别训练 GPR 模型。交叉验证采用 GroupKFold (k=5，按几何分组)，杜绝几何信息发生数据泄漏。")

    add_heading_2(doc, "2.3 仿真与预测结果对比")
    add_normal_paragraph(doc, "基于单质坡地数据集的测试表明，经典机器学习（XGBoost, GPR）在预测精度与效率上均全面完胜 CNN、LSTM、Transformer 与 DeepONet 等时序深度学习模型，且训练时间从小时级降低至秒级：")

    headers_accuracy = ["方案与模型分类", "全区平均误差 (MAE)", "峰值区平均误差 (MAE)", "全局判定系数 (R²)", "训练硬件与耗时说明"]
    data_accuracy = [
        ["改进方案 GPR (高斯过程)", "0.0244", "0.0218", "0.9474", "CPU 秒级训练，支持置信区间估计"],
        ["改进方案 XGBoost (树模型)", "0.0330", "0.0315", "0.9204", "CPU 秒级训练，支持特征重要性分析"],
        ["B2 Oracle 均值极限", "0.0469", "0.0480", "0.8285", "已知工况几何时的均值回归上限"],
        ["B3 规则网格插值", "0.0708", "0.0658", "0.6326", "非机器学习数学插值的精度下限"],
        ["B1 全局均值基线", "0.1030", "0.0767", "0.3214", "R² 零点盲猜基准"],
        ["初始时序 DeepONet", "0.0506", "—", "0.8354", "GPU 小时级训练，受地震波限制严重"],
        ["初始时序 CNN / LSTM", "0.0527 / 0.0544", "—", "0.8275 / 0.8147", "GPU 小时级训练，欠拟合严重"],
        ["初始时序 Transformer", "0.0572", "—", "0.7896", "GPU 小时级训练，小样本泛化极差"]
    ]
    create_styled_table(doc, headers_accuracy, data_accuracy)

    add_heading_2(doc, "2.4 诚实适用域外推分析")
    add_normal_paragraph(doc, "为评估模型的外推泛化边界，进行了单维度留一（Leave-One-Out）测试，结果表明代理模型“内插极强，外插易崩溃”：")
    
    headers_extrapolate = ["留一测试场景", "判定系数 (R²)", "平均绝对误差 (MAE)", "外推泛化性质与可信度结论"]
    data_extrapolate = [
        ["标准 GroupKFold (内插)", "0.946", "0.0249", "内插完好，适用于设计参数空间内插"],
        ["留一入射角外推", "0.924", "0.0290", "角度内插充分，响应连续稳健"],
        ["留一坡高 (h=50/100m)", "0.765 / 0.839", "0.0713 / 0.0648", "中等，处于插值范围内插可靠"],
        ["留一坡高 (h=200m)", "0.577", "0.0977", "外插临界边界，误差波动大，需慎用"],
        ["留一坡高 (h=10m 下界外插)", "−6.829", "0.0844", "崩溃失真，超出参数下限，模型失效"],
        ["留一坡高 (h=400m 上界外插)", "0.023", "0.1209", "崩溃失真，超出参数上限，模型失效"],
        ["留一地震波 (仅3条波外推)", "0.272", "0.1069", "波外推受限，主要因训练集波形太少"]
    ]
    create_styled_table(doc, headers_extrapolate, data_extrapolate)

    add_heading_2(doc, "2.5 机器学习下一步计划")
    add_bullet_paragraph(doc, "1. 并入多地层有限元数据：", "待 Part I 中 v8 新版脚本跑出的成层斜坡有限元工况数据计算完毕后，将多地层数据并入当前的数据集（从 522 样本扩充至 ~1000 样本）。")
    add_bullet_paragraph(doc, "2. 扩展输入地震波库：", "将地震波样本由目前的 3 条扩展至 10 条以上，以提取更丰富的谱物理特征，攻克留一波外推预测精度受限的问题。")
    add_bullet_paragraph(doc, "3. 研发多地层代理模型：", "基于无量纲参数组 η、d/h 和阻抗比 α，结合二维 POD 对 H(f, s) 曲面降维，研究多地层共振放大效应下的高维代理模型。")

    # ==================== 附录 ====================
    add_heading_1(doc, "附录")
    
    add_heading_2(doc, "A. 符号说明")
    headers_sym = ["符号", "含义说明"]
    data_sym = [
        ["TAF_h / TAF_v", "水平向 / 竖向地形放大系数（以基岩半空间自由场地表 PGA 归一）"],
        ["PGA_h / PGA_v", "水平向 / 竖向峰值地面加速度"],
        ["VAB", "粘弹性人工边界（Viscoelastic Artificial Boundary）"],
        ["POD / PCA", "本征正交分解（Proper Orthogonal Decomposition）"],
        ["GPR", "高斯过程回归（Gaussian Process Regression）"],
        ["MAE / R²", "平均绝对误差 / 决定系数"],
        ["CFL 守卫", "Courant-Friedrichs-Lewy 动力显式稳定性时步校验门槛"]
    ]
    create_styled_table(doc, headers_sym, data_sym)

    add_heading_2(doc, "B. 关键文件索引")
    headers_files = ["关键文件路径", "功能与技术细节描述"]
    data_files = [
        ["Modeling/Multi/VAB_oblique_TAF_multilayer_v8.py", "Abaqus 多地层有限元数值仿真主脚本（2479 行）"],
        ["Achievements/v2_to_v8_upgrade_report.md", "二维斜坡 SV 波斜入射有限元脚本从 v2 到 v8 升级对比报告"],
        ["Achievements/近期研究进展汇报.md", "近阶段研究进展 markdown 汇编文件"],
        ["ML/outputs_v3/", "基于均质单质坡地数据集（E:\\Abaqus）的经典机器学习 GPR 精度结果目录"]
    ]
    create_styled_table(doc, headers_files, data_files)

    # 保存
    output_path = r"c:\Users\12462\Documents\Code\AbqScripts\Achievements\近期研究进展汇报_v2_to_v8_建模详解版.docx"
    doc.save(output_path)
    print(f"Achievements report successfully updated: {output_path}")

if __name__ == "__main__":
    main()
